import os
import io
import re
import json
import zipfile
import difflib
import fitz
import docx
import pandas as pd
import google.generativeai as genai
import shutil 
from docx.shared import RGBColor
import PIL.Image
from flask import (Flask, request, jsonify, render_template, send_file, make_response, redirect, url_for, flash)
from dotenv import load_dotenv
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Inches
from docx import Document 
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import or_
from flask_login import (LoginManager, UserMixin, login_user, logout_user, login_required, current_user)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import datetime 
load_dotenv()

RPM_LIMIT = 10 
gemini_usage_count = 0
gemini_last_reset = datetime.datetime.now()

def track_gemini_usage():
    global gemini_usage_count, gemini_last_reset
    
    now = datetime.datetime.now()
    time_diff = (now - gemini_last_reset).total_seconds()

    if time_diff >= 60:
        gemini_usage_count = 0
        gemini_last_reset = now
    
    gemini_usage_count += 1

app = Flask(__name__)

APP_ROOT = os.path.dirname(os.path.abspath(__file__))
HF_DATA_DIR = '/data' 

if os.path.exists(HF_DATA_DIR):
    print("Menggunakan Persistent Storage Hugging Face (/data)")
    
    db_path = os.path.join(HF_DATA_DIR, 'database.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
    initial_db_source = os.path.join(APP_ROOT, 'instance', 'database.db')
    if not os.path.exists(db_path) and os.path.exists(initial_db_source):
        print(f"Menyalin database awal dari {initial_db_source} ke {db_path}...")
        try:
            shutil.copy2(initial_db_source, db_path)
            os.chmod(db_path, 0o666) 
            print("Database awal berhasil disalin.")
        except Exception as e:
            print(f"Gagal menyalin database awal: {e}")
    app.config['UPLOAD_FOLDER'] = os.path.join(HF_DATA_DIR, 'uploads')
    
else:
    print("--- Menggunakan Storage Lokal (Development) ---")
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///instance/database.db'
    app.config['UPLOAD_FOLDER'] = os.path.join(APP_ROOT, 'data')

app.config['SECRET_KEY'] = 'kunci-rahasia-anda-yang-acak-dan-kuat'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024 

APP_ROOT = os.path.dirname(os.path.abspath(__file__))

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login' 
login_manager.login_message = 'Silakan login untuk mengakses halaman ini.'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

try:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY tidak ditemukan di file .env")
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash') 
except Exception as e:
    print(f"Error saat mengkonfigurasi Google AI: {e}")

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    fullname = db.Column(db.String(200), nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    label = db.Column(db.String(100), nullable=False)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class AnalysisLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    feature_type = db.Column(db.String(50), nullable=False)
    document_type = db.Column(db.String(100), nullable=True)
    start_time = db.Column(db.DateTime, nullable=False, default=datetime.datetime.utcnow)
    end_time = db.Column(db.DateTime, nullable=True)
    deadline = db.Column(db.DateTime, nullable=True) 
    status = db.Column(db.String(20), nullable=False, default='unfinished') 
    user = db.relationship('User', backref='analysis_logs')

class AmsMonitoring(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    session_name = db.Column(db.String(200), nullable=False)
    monitoring_type = db.Column(db.String(50), nullable=False, default='standard') 
    periode = db.Column(db.Date, nullable=False, default=datetime.date.today)
    auditee = db.Column(db.String(200), nullable=False)
    tahun_audit = db.Column(db.Integer, nullable=False)
    total_rekomendasi = db.Column(db.Integer, nullable=False, default=0)
    selesai = db.Column(db.Integer, nullable=False, default=0)
    tidak_selesai = db.Column(db.Integer, nullable=False, default=0) 
    todo = db.Column(db.Integer, nullable=False, default=0) 
    belum_sesuai = db.Column(db.Integer, nullable=False, default=0) 
    belum_tl = db.Column(db.Integer, nullable=False, default=0) 
    tdd = db.Column(db.Integer, nullable=False, default=0) 
    user = db.relationship('User', backref='ams_monitorings')

class AmsReminder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    assigned_to = db.Column(db.String(200), nullable=False)
    pic_auditee = db.Column(db.Text, nullable=True) 
    pic_skai = db.Column(db.Text, nullable=True)
    pic_reminder = db.Column(db.Text, nullable=True)
    temuan = db.Column(db.Text, nullable=True)    
    subject = db.Column(db.String(255), nullable=False)
    deadline = db.Column(db.Date, nullable=False)
    status = db.Column(db.String(50), default='On Progress')
    is_responded = db.Column(db.Boolean, default=False, nullable=False)
    is_reminded = db.Column(db.Boolean, default=False, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.now)
    linked_temuan_ids = db.Column(db.Text, nullable=True)
    linked_by = db.Column(db.String(100), nullable=True)
    user = db.relationship('User', backref='ams_reminders')
    tembusan = db.Column(db.String(200), nullable=True)
    
class SharedMonitoring(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    session_name = db.Column(db.String(200), nullable=False)
    owner = db.relationship('User', foreign_keys=[owner_id], backref='owned_monitoring_shares')
    shared_with = db.relationship('User', foreign_keys=[shared_with_id], backref='received_monitoring_shares')

class SharedTemuanSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    session_id = db.Column(db.Integer, db.ForeignKey('ams_temuan_session.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    session = db.relationship('AmsTemuanSession', backref='shares')
    user = db.relationship('User', foreign_keys=[owner_id])

class AmsTlTidakSetuju(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    no_aoi = db.Column(db.String(100))
    jenis_aoi = db.Column(db.String(100))
    klasifikasi = db.Column(db.String(100))
    no_lha = db.Column(db.String(100))
    nama_penugasan = db.Column(db.String(200))
    keterangan = db.Column(db.Text)
    temuan = db.Column(db.Text)
    rekomendasi = db.Column(db.Text)
    auditee = db.Column(db.String(200))
    target_per_lha = db.Column(db.Date, nullable=True)
    perubahan_target_date = db.Column(db.Date, nullable=True)
    tindak_lanjut = db.Column(db.Text)
    user = db.relationship('User', backref='ams_tl_tidak_setuju')

class SharedFolder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    folder_name = db.Column(db.String(200), nullable=False)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    folder_type = db.Column(db.String(50), nullable=False, default='general') 
    owner = db.relationship('User', foreign_keys=[owner_id], backref='owned_shares')
    shared_with_user = db.relationship('User', foreign_keys=[shared_with_id], backref='received_shares')

class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False) 
    folder_name = db.Column(db.String(200), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    row_id = db.Column(db.Integer, nullable=False) 
    username = db.Column(db.String(100), nullable=False) 
    text = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    parent_id = db.Column(db.Integer, db.ForeignKey('comment.id'), nullable=True)
    replies = db.relationship('Comment', backref=db.backref('parent', remote_side=[id]), lazy='dynamic')
    __table_args__ = (db.UniqueConstraint('owner_id', 'folder_name', 'filename', 'row_id', name='_unique_comment_target'),)

class RowAction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    folder_name = db.Column(db.String(200), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    row_id = db.Column(db.Integer, nullable=False)
    is_ganti = db.Column(db.Boolean, default=False, nullable=False)
    pic_user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    pic_user = db.relationship('User', foreign_keys=[pic_user_id])
    __table_args__ = (db.UniqueConstraint('owner_id', 'folder_name', 'filename', 'row_id', name='_unique_row_action'),)

class AmsRowHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    row_id = db.Column(db.Integer, db.ForeignKey('ams_temuan_row.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    username = db.Column(db.String(100), nullable=False)
    changed_columns = db.Column(db.Text, nullable=False) 
    timestamp = db.Column(db.DateTime, default=datetime.datetime.now)
    row = db.relationship('AmsTemuanRow', backref=db.backref('history', cascade="all, delete-orphan"))

class AmsTemuanSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    nama_sesi = db.Column(db.String(200), nullable=False)
    jenis_audit = db.Column(db.String(100), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    rows = db.relationship('AmsTemuanRow', backref='session', cascade="all, delete-orphan", lazy=True)
    user = db.relationship('User', backref='ams_temuan_sessions')

class SharedTlTidakSetuju(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class AmsTemuanRow(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    session_id = db.Column(db.Integer, db.ForeignKey('ams_temuan_session.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    pic_skai = db.Column(db.String(200), nullable=True)
    no_aoi = db.Column(db.String(100))
    jenis_aoi = db.Column(db.String(100))
    klasifikasi = db.Column(db.String(100))
    no_lha = db.Column(db.String(100))
    nama_penugasan = db.Column(db.String(200))
    aoi = db.Column(db.Text) 
    rekomendasi = db.Column(db.Text)
    rencana_tl = db.Column(db.Text)
    rencana_evidence = db.Column(db.Text)
    auditee = db.Column(db.String(200))
    pic_auditee = db.Column(db.Text) 
    target_penyelesaian = db.Column(db.String(50), nullable=True)
    perubahan_target = db.Column(db.String(50), nullable=True)
    tindak_lanjut = db.Column(db.Text)
    signifikansi = db.Column(db.String(100))
    jml_rekomendasi = db.Column(db.Integer, default=0)
    selesai = db.Column(db.Integer, default=0)
    belum_jt_bs = db.Column(db.Integer, default=0) 
    os_bd = db.Column(db.Integer, default=0) 
    tdd = db.Column(db.Integer, default=0) #
    control = db.Column(db.String(200))
    auditor_notes = db.Column(db.Text, nullable=True)
    last_modified_by = db.Column(db.String(100), nullable=True)
    last_modified_at = db.Column(db.DateTime, nullable=True)

class AmsTemuanComment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    row_id = db.Column(db.Integer, db.ForeignKey('ams_temuan_row.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    username = db.Column(db.String(100))
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.now)
    row = db.relationship('AmsTemuanRow', backref=db.backref('comments', cascade="all, delete-orphan"))

class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sender_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    recipient_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    subject = db.Column(db.String(255), nullable=False)
    body = db.Column(db.Text, nullable=True)
    attachment_path = db.Column(db.String(500), nullable=True) 
    original_filename = db.Column(db.String(255), nullable=True) 
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.datetime.utcnow)
    is_read = db.Column(db.Boolean, default=False, nullable=False)
    sender = db.relationship('User', foreign_keys=[sender_id], backref='sent_messages')
    recipient = db.relationship('User', foreign_keys=[recipient_id], backref='received_messages')

class LibraryFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shelf_id = db.Column(db.String(50), nullable=True)
    title = db.Column(db.String(255), nullable=False)
    category = db.Column(db.String(100), nullable=False)
    summary = db.Column(db.Text, nullable=True)
    cluster = db.Column(db.String(200), nullable=True)
    filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_size = db.Column(db.String(50), nullable=False)
    file_type = db.Column(db.String(20), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.datetime.now)
    user = db.relationship('User', backref='library_files')

class AmsAuditor(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    department = db.Column(db.String(100), nullable=True)
    total_temuan = db.Column(db.Integer, default=0)
    selesai = db.Column(db.Integer, default=0)
    bjt = db.Column(db.Integer, default=0)
    outstanding = db.Column(db.Integer, default=0)
    period = db.Column(db.String(7), nullable=False) # Format: YYYY-MM
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    user = db.relationship('User', backref='ams_auditors')

@login_manager.user_loader
def load_user(user_id):
    """Fungsi wajib untuk Flask-Login"""
    return User.query.get(int(user_id))

@app.cli.command("create-db")
def create_db_command():
    with app.app_context():
        db.create_all() 
        
        users_data = [
            ("Deny", "Deny Syahbani", "KangDenys#7601", "CoE"),
            ("Fadian", "Fadian Dwiantara", "BosFadianD#5710", "CoE"),
            ("Winda", "Winda Anggraeni", "WindAng#1956", "CoE"),
            ("Renzie", "Renzie Aditya", "Renzie1234", "CoE"),
            ("Hari", "Hari Sundoro", "MrHariSund#2750", "Advisory"),
            ("Rizky", "Rizky Ananda Putra", "KingofJokes#1963", "Advisory"),
            ("Rosyid", "M. Rosyid Ridho Muttaqien", "LordOcidAI#1706", "Audit"),
            ("Made", "I Made Suandi Putra", "IMadeSP#1425", "Audit"),
            ("Wira", "Wirawan Arief", "MisterWAN#9661", "Advisory"),
            ("Darmo", "Darmo Saputro Wibowo", "MrDSW#8814", "Audit"),
            ("Jaka", "Jaka Tirtana Hanafiah", "KingJakaTH#6197", "Audit"),
            ("Lucky", "Lucky Parwitasari", "MrsLucky#4963", "Audit"),
            ("Handaru", "Handarudigdaya J.K.", "SirDaruJK#0175", "Audit"),
            ("Dhita", "Aliya Anindhita Rachman", "ADhitaR#8513", "Audit"),
            ("Uni", "Laila Fajriani", "UniLaila#9265", "Advisory"),
            ("Bakhas", "Bakhas Nasrani Diso", "DisoKingIT#1760", "Advisory"),
            ("Jihan", "Jihan Abigail", "JihanCoE#1850", "CoE"),
            ("Fajar", "Fajar Setianto", "BangFajarS#5610", "Audit")
        ]
        
        for username, fullname, password, label in users_data:
            if not User.query.filter_by(username=username).first():
                user = User(username=username, fullname=fullname, label=label)
                user.set_password(password)
                db.session.add(user)
                print(f"User {username} ({fullname}) dibuat.")
        
        db.session.commit()
        print("Database dan user awal telah selesai dibuat.")

def get_user_root_folder():
    """Mendapatkan path folder root pengguna saat ini."""
    if not current_user.is_authenticated:
        return None
    user_id_str = str(current_user.id)
    if not user_id_str.isalnum():
        raise ValueError("User ID tidak valid untuk path folder.")     
    user_folder_path = os.path.join(app.config['UPLOAD_FOLDER'], user_id_str)
    os.makedirs(user_folder_path, exist_ok=True)
    return user_folder_path

def create_user_folder(folder_name):
    """Membuat sub-folder di folder root pengguna."""
    root_folder = get_user_root_folder()
    if not root_folder:
        raise Exception("Pengguna tidak terautentikasi.")
    clean_folder_name = re.sub(r'[^\w\s-]', '', folder_name).strip() 
    clean_folder_name = re.sub(r'[-\s]+', '_', clean_folder_name) 
    
    if not clean_folder_name:
        raise ValueError("Nama folder tidak valid setelah dibersihkan.")

    new_folder_path = os.path.join(root_folder, clean_folder_name)
    if os.path.exists(new_folder_path):
        raise ValueError("Folder dengan nama yang sama sudah ada.")

    os.makedirs(new_folder_path)
    return clean_folder_name

def get_user_folders():
    """
    MODIFIKASI: Mengambil folder milik sendiri DAN folder yang di-share.
    Mengembalikan list of dictionaries.
    """
    root_folder = get_user_root_folder()
    if not root_folder:
        return []
    
    my_folders = []
    try:
        owned_folders = [
            d for d in os.listdir(root_folder) 
            if os.path.isdir(os.path.join(root_folder, d))
        ]
        for folder in owned_folders:
            my_folders.append({
                "name": folder,
                "owner_name": current_user.username,
                "is_owner": True,
                "owner_id": current_user.id
            })
    except Exception as e:
        print(f"Error saat listing folder di {root_folder}: {e}")
    shared_folders = SharedFolder.query.filter_by(shared_with_id=current_user.id).all()
    
    for share in shared_folders:
        owner_root_path = os.path.join(app.config['UPLOAD_FOLDER'], str(share.owner_id))
        folder_path = os.path.join(owner_root_path, share.folder_name)
        
        if os.path.isdir(folder_path):
            my_folders.append({
                "name": share.folder_name,
                "owner_name": share.owner.username,
                "is_owner": False,
                "owner_id": share.owner_id 
            })

    return my_folders

def _extract_text_with_pages(file_bytes, file_extension):
    pages_content = []
    if file_extension == 'pdf':
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num, page in enumerate(pdf_document):
                pages_content.append({"halaman": page_num + 1, "teks": page.get_text()})
            pdf_document.close()
        except Exception as e:
            raise ValueError(f"Gagal membaca file PDF: {e}")
    elif file_extension == 'docx':
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            pages_content.append({"halaman": 1, "teks": full_text})
        except Exception as e:
            raise ValueError(f"Gagal membaca file DOCX: {e}")
    else:
        raise ValueError("Format file tidak didukung. Harap unggah .pdf atau .docx")
    return pages_content

def _get_text_from_flask_file(file):
    file_bytes = file.read()
    file.seek(0) 
    file_extension = file.filename.split('.')[-1].lower()
    return _extract_text_with_pages(file_bytes, file_extension)

def extract_sentences_with_pages(pages_content):
    sentences_with_pages = []
    for page in pages_content:
        page_num = page['halaman']
        text = page['teks']
        sentences = split_text_into_sentences(text) 
        for sentence in sentences:
            sentences_with_pages.append({'sentence': sentence, 'page': page_num})
    return sentences_with_pages

def _get_full_text_from_file(file):
    file_bytes = file.read()
    file.seek(0)
    file_extension = file.filename.split('.')[-1].lower()
    
    pages = _extract_text_with_pages(file_bytes, file_extension)
    return "\n".join([page['teks'] for page in pages])

def proofread_with_gemini(text_to_check):
    if not text_to_check or text_to_check.isspace():
        return []
    prompt = f"""
    Anda adalah seorang auditor dan ahli bahasa Indonesia yang sangat teliti. Anda diberikan dokumen dan tugas Anda adalah melakukan proofread pada teks berikut. Fokus pada:
    1. Memperbaiki kesalahan ketik (typo) agar semuanya sesuai dengan standar KBBI dan PUEBI.
    1. Kalau ada kata-kata yang tidak sesuai KBBI dan PUEBI, tolong jangan highlight semua kalimatnya, tapi cukup highlight kata-kata yang salah serta perbaiki kata-kata itu aja, jangan perbaiki semua kalimatnya
    3. Jika ada kata yang diitalic, biarkan saja
    4. Nama-nama yang diberi ini pastikan benar juga "Yullyan, I Made Suandi Putra, Laila Fajriani, Hari Sundoro, Bakhas Nasrani Diso, Rizky Ananda Putra, Wirawan Arief Nugroho, Lelya Novita Kusumawati, Ryani Ariesti Syafitri, Darmo Saputro Wibowo, Lucky Parwitasari, Handarudigdaya Jalanidhi Kuncaratrah, Fajar Setianto, Jaka Tirtana Hanafiah, tMuhammad Rosyid Ridho Muttaqien, Octovian Abrianto, Deny Sjahbani, Jihan Abigail, Winda Anggraini, Fadian Dwiantara, Aliya Anindhita Rachman"
    5. Fontnya arial dan jangan diganti. Khusus untuk judul paling atas, itu font sizenya 12 dan bodynya selalu 11
    6. Khusus "Indonesia Financial Group (IFG)", meskipun bahasa inggris, tidak perlu di italic
    7. Kalau ada kata yang sudah diberikan akronimnya di awal, maka di halaman berikut-berikutnya cukup akronimnya saja, tidak perlu ditulis lengkap lagi
    8. Pada bagian Nomor surat dan Penutup tidak perlu dicek, biarkan seperti itu
    9. Ketika Anda perbaiki, fontnya pastikan Arial dengan ukuran 11 juga (Tidak diganti)
    10. Pada kalimat "Indonesia Financial Group", jika terdapat kata typo "Finansial", tolong Anda sarankan untuk ganti ke "Financial"
    11. Yang benar adalah "Satuan Kerja Audit Internal", bukan "Satuan Pengendali Internal Audit"
    12. Jika terdapat kata "reviu", biarkan itu sebagai benar
    13. Kalau ada kata "IM", "ST", "SKAI", "IFG", "TV (Angka Romawi)", "RKAT", dan "RKAP", itu tidak perlu ditandai sebagai salah dan tidak perlu disarankan untuk italic / bold / underline
    14. Untuk nama modul seperti "Modul Sourcing, dll", itu tidak perlu italic
    15. Kalau ada kata dalam bahasa inggris yang masih masuk akal dan nyambung dengan kalimat yang dibahas, tidak perlu Anda sarankan untuk ganti ke bahasa indonesia
    16. Jika ada bahasa inggris dan akronimnya seperti "General Ledger (GL)", tolong dilakukan italic pada kata tersebut pada saat download file hasil revisinya, akronimnya tidak perlu diitalic
    17. Awal kalimat selalu dimulai dengan huruf kapital. Jika akhir poin diberi tanda ";", maka poin selanjutnya tidak perlu kapital
    18. Di file hasil revisi, Anda jangan ganti dari yang aslinya. Misalnya kalau ada kata yang diitalic di file asli, jangan Anda hilangkan italicnya
    19. Tolong perhatikan juga tanda bacanya, seperti koma, titik koma, titik, tanda hubung, dan lain-lain. Pastikan sesuai dan ada tanda titik di setiap akhir kalimat
    20. Kalau ada bahasa inggris yang belum diitalic, tolong diitalic
    21. Kata Internal Memorandum itu tidak perlu diitalic karena itu nama dari sebuah dokumen

    PENTING: Berikan hasil dalam format yang SANGAT KETAT. Untuk setiap kesalahan, gunakan format:
    [SALAH] kata atau frasa yang salah -> [BENAR] kata atau frasa perbaikan -> [KALIMAT] kalimat lengkap asli tempat kesalahan ditemukan

    Contoh:
    [SALAH] dikarenakan -> [BENAR] karena -> [KALIMAT] Hal itu terjadi dikarenakan kelalaian petugas.

    Jika tidak ada kesalahan sama sekali, kembalikan teks: "TIDAK ADA KESALAHAN"

    Berikut adalah teks yang harus Anda periksa:
    ---
    {text_to_check}
    """
    try:
        track_gemini_usage()
        response = model.generate_content(prompt)
        pattern = re.compile(r"\[SALAH\]\s*(.*?)\s*->\s*\[BENAR\]\s*(.*?)\s*->\s*\[KALIMAT\]\s*(.*?)\s*(\n|$)", re.IGNORECASE | re.DOTALL)
        found_errors = pattern.findall(response.text)
        return [{"salah": salah.strip(), "benar": benar.strip(), "kalimat": kalimat.strip()} for salah, benar, kalimat, _ in found_errors]
    except Exception as e:
        print(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return [{"salah": "ERROR", "benar": str(e), "kalimat": "Gagal menghubungi API"}]

def split_text_into_sentences(full_text):
    """
    Memecah teks penuh menjadi daftar kalimat.
    Ini adalah implementasi sederhana dan mungkin tidak sempurna untuk semua kasus,
    tetapi merupakan awal yang baik untuk dokumen formal.
    """
    if not full_text:
        return []

    sentences = re.split(r'(?<=[.!?])\s+', full_text)
    return [s.strip() for s in sentences if len(s.strip()) > 10]

def analyze_document_by_section(original_text, revised_text):
    if not original_text or not revised_text:
        return []
    prompt = f"""
    Anda adalah seorang auditor ahli. Tugas Anda adalah membandingkan dua dokumen: Dokumen Asli dan Dokumen Revisi.
    Tujuannya adalah untuk memastikan bahwa konten di setiap sub-bab pada Dokumen Revisi tetap sejalan dengan makna dan konteks sub-bab yang sesuai di Dokumen Asli.

    Instruksi:
    1.  Baca dan pahami struktur kedua dokumen. Identifikasi semua bab dan sub-bab.
    2.  ABAIKAN bagian tabel, gambar, daftar pustaka, dan lampiran.
    3.  Fokus analisis pada isi paragraf dan poin-poin.
    4.  Untuk setiap sub-bab di Dokumen Revisi, cari padanannya (referensinya) di Dokumen Asli.
    5.  Identifikasi kalimat di Dokumen Revisi yang memiliki perubahan makna signifikan.
    6.  Jelaskan alasannya dengan struktur ketat berikut:
        - Poin 1: Makna di dokumen asli (Jelaskan makna di dokumen asli).
        - Poin 2: Makna di dokumen yang dibanding (Jelaskan makna/konteks di dokumen revisi yang menyimpang).
        - Rekomendasi: (Saran perbaikan konkret).

    Berikut adalah kedua dokumen tersebut:
    ---
    DOKUMEN ASLI:
    {original_text}
    ---
    DOKUMEN REVISI:
    {revised_text}
    ---

    Berikan hasil analisis dalam format JSON array. Setiap objek HARUS memiliki 4 key berikut:
    1.  "sub_bab_asal": Nama sub-bab (dari Dokumen Revisi).
    2.  "sub_bab_referensi": Nama sub-bab (dari Dokumen Asli) yang menjadi referensi/padanan.
    3.  "kalimat_menyimpang": Kalimat dari Dokumen Revisi yang maknanya menyimpang.
    4.  "alasan": Gabungan teks Poin 1, Poin 2, dan Rekomendasi (Gunakan format: "1. Makna di dokumen asli: ... \\n2. Makna di dokumen yang dibanding: ... \\nRekomendasi: ...").

    Contoh Output:
    [
      {{
        "sub_bab_asal": "1.2 Tujuan Audit",
        "sub_bab_referensi": "Bab I: Pendahuluan - Tujuan",
        "kalimat_menyimpang": "Manajemen wajib memastikan kepatuhan...",
        "alasan": "1. Makna di dokumen asli: Kewajiban hanya pada hukum nasional.\\n2. Makna di dokumen yang dibanding: Menambahkan hukum internasional yang tidak relevan.\\nRekomendasi: Hapus frasa 'hukum internasional'."
      }}
    ]

    PENTING: HANYA KELUARKAN JSON ARRAY MURNI.
    """
    try:
        track_gemini_usage()
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        response_text = re.sub(r'```json\s*|\s*```', '', response_text)
        analysis_result = json.loads(response_text)
        if not isinstance(analysis_result, list):
            raise ValueError("Respons dari AI bukanlah sebuah array JSON.")
        return analysis_result
    except Exception as e:
        print(f"Error analysis: {e}")
        return [{"sub_bab_asal": "Error", "sub_bab_referensi": "-", "kalimat_menyimpang": "Gagal memproses respons AI.", "alasan": str(e)}]


def analyze_context_difference(original_sentence, revised_sentence):
    """
    Menganalisis mengapa konteks kalimat revisi berbeda dari kalimat asli.
    Versi ini lebih sederhana karena nomor halaman disediakan oleh backend.
    """
    if not original_sentence or not revised_sentence:
        return {"alasan": "Tidak cukup data untuk dianalisis.", "kalimat_menyimpang": revised_sentence}

    prompt = f"""
    Anda adalah seorang auditor ahli. Anda diberikan dua dokumen, satu dokumen asli dan satu dokumen lainnya
    Tugas Anda adalah menganalisis mengapa konteks, makna, fokus kalimat / paragraf pada dokumen revisi berbeda dari dokumen asli. Beberapa ketentuannya sebagai berikut
    1. Struktur dari dokumen asli dengan dokumen revisi berbeda, tetapi membahas mengenai hal yang sama.
    2. Anda pelajari terlebih dahulu di dokumen Asli itu maknanya seperti apa secara detail terutama setiap poin poinnya, begitupun juga untuk di dokumen revisi
    3. Tolong Anda periksa secara detail apakah ada perubahan makna yang signifikan pada kalimat di dokumen revisi dibandingkan dokumen aslinya. 
    4. Identifikasi secara teliti semua kalimat yang ada di dokumen revisi yang memiliki perubahan makna signifikan dibandingkan dokumen aslinya.
    5. Berikan penjelasan singkat dan jelas mengapa kalimat tersebut diubah (misalnya: "Menambahkan detail spesifik", "Mengubah fokus dari A ke B", "Memperjelas ambigu", "Mengoreksi fakta").
    6. Jelaskan konteks atau makna utama kalimat asli dan kalimat revisi secara singkat.
    7. Tampilkan alasan yang sangat detail karena ini adalah dokumen Audit sehingga harus sangat teliti dan membuat para Auditor paham mengapa ada perubahan makna pada kalimat tersebut.
    8. Yang ditempilkan pada tabel hasil itu nanti jangan seluruh paragrafnya, tapi cukup kalimat yang mengalami perubahan makna signifikan saja. 

    Kalimat Asli: "{original_sentence}"
    Kalimat Revisi: "{revised_sentence}"

    Analisis dan identifikasi apakah terjadi perubahan makna yang signifikan. Jelaskan alasan perubahannya secara ringkas namun jelas.

    Berikan analisis Anda secara eksklusif dalam format JSON objek dengan dua kunci berikut:
    1. "alasan": Jelaskan secara ringkas mengapa kalimat tersebut diubah. Jika tidak ada perubahan makna signifikan, isi dengan "Tidak ada perubahan makna signifikan."
    2. "kalimat_menyimpang": Salin kembali kalimat revisi yang mengalami perubahan makna.

    Contoh Output:
    {{
        "alasan": "Menambahkan kewajiban hukum yang spesifik.",
        "kalimat_menyimpang": "Manajemen harus memastikan kepatuhan terhadap peraturan perundang-undangan yang berlaku."
    }}

    PENTING: HANYA KELUARKAN OBJEK JSON MURNI. TANPA TEKS PENDAHULU ATAU PENUTUP.
    """

    try:
        track_gemini_usage()
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        try:
            analysis_result = json.loads(response_text)
        except json.JSONDecodeError:
            print(f"[DEBUG] Gagal parsing JSON langsung. Mencoba ekstraksi. Response: {response_text}")
            match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if match:
                json_string = match.group(0)
                analysis_result = json.loads(json_string)
            else:
                raise ValueError("Tidak dapat menemukan objek JSON yang valid dalam respons AI.")

        required_keys = ["alasan", "kalimat_menyimpang"]
        if not all(key in analysis_result for key in required_keys):
            print(f"[DEBUG] Struktur JSON tidak lengkap. Diterima: {analysis_result}")
            raise ValueError("Struktur JSON dari AI tidak lengkap atau tidak sesuai format.")
            
        return analysis_result

    except Exception as e:
        print(f"[ERROR] Terjadi kesalahan saat menganalisis konteks: {e}")
        print(f"[ERROR] Original: {original_sentence}")
        print(f"[ERROR] Revised: {revised_sentence}")
        print(f"[ERROR] AI Response was: {response.text if 'response' in locals() else 'N/A'}")
        
        return {
            "alasan": f"Error: AI gagal memberikan respons yang valid. ({str(e)})",
            "kalimat_menyimpang": revised_sentence
        }

def analyze_document_coherence(full_text):
    if not full_text or full_text.isspace():
        return []

    prompt = f"""
    Anda adalah seorang auditor ahli yang bertugas menganalisis struktur dan koherensi sebuah tulisan.
    Tugas Anda adalah membaca keseluruhan teks berikut dan mengidentifikasi setiap kalimat atau paragraf yang tidak koheren atau keluar dari topik utama di dalam sebuah sub-bagian.
    
    Untuk setiap ketidaksesuaian yang Anda temukan, lakukan hal berikut:
    1. Bacalah mengenai judul dari section atau subsection yang ada pada file tersebut, terutama bacalah isi paragrafnya dan makna setiap kalimatnya
    2. Tentukan topik utama dari setiap section / subsection terutama isi paragrafnya.
    3. Identifikasi kalimat asli yang menyimpang dari topik tersebut.
    4. Berikan saran dengan menghighlight kalimat tersebut untuk ditulis ulang (rewording) agar relevan dan menyatu kembali dengan topik utamanya, ikuti standar KBBI, PUEBI, dan SPOK Bahasa Indonesia.
    5. Jika Anda memiliki asumsi atau catatan tambahan tentang revisi tersebut (seperti "asumsi logis" atau "catatan: ..."), PISAHKAN catatan itu.

    Berikan hasil dalam format yang SANGAT KETAT.
    Format 1 (Tanpa Catatan):
    [TOPIK UTAMA] topik utama -> [TEKS ASLI] kalimat asli -> [SARAN REVISI] versi kalimat yang sudah diperbaiki

    Format 2 (Dengan Catatan):
    [TOPIK UTAMA] topik utama -> [TEKS ASLI] kalimat asli -> [SARAN REVISI] versi kalimat yang sudah diperbaiki -> [CATATAN] asumsi atau catatan Anda

    Contoh Format 2:
    [TOPIK UTAMA] Rencana Kerja Tahunan -> [TEKS ASLI] Penyebab utamanya adalah... -> [SARAN REVISI] Penyebab utamanya adalah... -> [CATATAN] Asumsi logis dari konteks.

    Jika seluruh dokumen sudah koheren dan tidak ada masalah, kembalikan teks: "TIDAK ADA MASALAH KOHERENSI"

    Teks:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(
            r"\[TOPIK UTAMA\]\s*(.*?)\s*->\s*\[TEKS ASLI\]\s*(.*?)\s*->\s*\[SARAN REVISI\]\s*(.*?)\s*(?:->\s*\[CATATAN\]\s*(.*?)\s*)?(\n|$)", 
            re.IGNORECASE | re.DOTALL
        )
        found_issues = pattern.findall(response.text)
        
        results = []
        for topik, asli, saran, catatan, _ in found_issues:
            results.append({
                "topik": topik.strip(), 
                "asli": asli.strip(), 
                "saran": saran.strip(), 
                "catatan": catatan.strip() if catatan else ""
            })
        return results
        
    except Exception as e:
        print(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return [{"topik": "ERROR", "asli": str(e), "saran": "Gagal menghubungi API", "catatan": ""}]

def get_structural_recommendations(full_text):
    if not full_text or full_text.isspace():
        return []
    prompt = f"""
    Anda adalah seorang auditor ahli yang bertugas untuk melakukan analisis terhadap dokumen. Tugas Anda adalah menemukan paragraf yang terkesan 'salah tempat' dan memberikan saran di bagian mana seharusnya paragraf tersebut berada saat ini (lokasi asli).

    Untuk setiap paragraf yang terdeteksi, Anda harus:
    1. Bacalah semua dokumennya terlebih dahulu, temukan ide-ide utama di setiap paragraf, dan merevisi jika perlu.
    2. Pada saat Anda membaca dokumennya, tolong identifikasi teks lengkap dari paragraf yang berada tidak pada tempatnya.
    3. Tentukan di sub-bab atau sub-bab mana paragraf itu berada saat ini (lokasi asli).
    4. Berikan rekomendasi di sub-bab atau sub-bab mana paragraf tersebut seharusnya diletakkan agar lebih koheren dan logis.
    5. Kalau ada bagian yang harus dipindahkan ke Ringkasan Eksekutif, itu tidak perlu dimasukkan ke dalam tabel.
    6. Kalau ada kata yang merupakan bahasa inggris, biarkan saja dan tidak perlu diitalic.
    7. Kalau ada kata yang tidak baku sesuai dengan standar KBBI, harap Anda perbaiki saja. Sehingga kata tersebut menjadi kata baku.
    8. Pada bagian lampiran, tidak perlu di cek/dicek untuk dipindahkan ke bagian lainnya karena itu sudah fixed / sudah benar

    Berikan hasil dalam format JSON yang berisi sebuah list. Setiap objek harus memiliki tiga kunci: "misplaced_paragraph", "original_section", dan "recommended_section".

    Contoh Format JSON:
    [
      {{
        "misplaced_paragraph": "Selain itu, audit internal juga bertugas memeriksa laporan keuangan setiap kuartal...",
        "original_section": "Bab 2.1: Prosedur Whistleblowing",
        "recommended_section": "Bab 4.2: Peran Audit Internal"
      }}
    ]

    Berikut adalah teks yang harus Anda periksa:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        cleaned_response = re.sub(r'[—–]', '-', response.text.strip()) 
        cleaned_response = re.sub(r'```json\s*|\s*```', '', cleaned_response)
        return json.loads(cleaned_response)
    except Exception as e:
        print(f"Failed to Generate Response from AI: {e}")
        return [{"misplaced_paragraph": "Error: " + str(e), "original_section": "Gagal menghubungi API", "recommended_section": "Periksa prompt Anda."}]

def review_document_comprehensive(text_to_check):
    if not text_to_check or text_to_check.isspace():
        return []

    prompt = f"""
    Anda adalah Editor Kepala dan Auditor Senior. Tugas Anda adalah melakukan "Reviu Dokumen Lengkap" pada teks berikut.
    Lakukan 3 jenis analisis sekaligus secara mendalam:

    1. **Proofreading (Typo & PUEBI):** Cari kesalahan ejaan, tanda baca, dan kata tidak baku sesuai KBBI. Kalau ada yang bahasa inggris, italic aja, jangan diganti ke bahasa Indonesia, nomor surat gausah dicek, kalau sudah dalam bahasa indonesia, gausah ditranslate ke bahasa inggris
    2. **Koherensi (Rewording):** Identifikasi kalimat yang tidak nyambung dengan topik utama paragraf atau sulit dipahami.
    3. **Struktur (Restrukturisasi):** Identifikasi paragraf yang salah tempat atau tidak logis urutannya.

    Teks untuk dianalisis:
    ---
    {text_to_check}
    ---

    **FORMAT OUTPUT:**
    Berikan jawaban HANYA dalam format JSON Array murni. Setiap objek harus memiliki kunci berikut (pastikan nama kunci persis sama):
    - "kategori": Isi dengan salah satu: "Proofreading", "Koherensi", atau "Struktur".
    - "masalah": Teks asli yang bermasalah (kata/kalimat/paragraf).
    - "saran": Perbaikan yang disarankan (kata baku/kalimat revisi/lokasi baru).
    - "penjelasan": Alasan singkat kenapa ini salah atau saran perbaikan.
    - "lokasi": Konteks lokasi (misal: "Paragraf 1", "Kalimat ke-2").
    """
    try:
        track_gemini_usage()
        response = model.generate_content(prompt)
        raw_response = response.text.strip()
        cleaned_text = re.sub(r'^```json\s*', '', raw_response)
        cleaned_text = re.sub(r'^```\s*', '', cleaned_text)
        cleaned_text = re.sub(r'\s*```$', '', cleaned_text)
        results = json.loads(cleaned_text)
        priority_map = {'Proofreading': 1, 'Koherensi': 2, 'Struktur': 3}
        results.sort(key=lambda x: priority_map.get(x.get('kategori'), 99))
        
        return results

    except Exception as e:
        print(f"Error Reviu Dokumen: {e}")
        return [{"kategori": "Error", "masalah": "Gagal analisis", "saran": "-", "penjelasan": str(e), "lokasi": "-"}]

def generate_revised_docx(file_bytes, errors):
    doc = docx.Document(io.BytesIO(file_bytes))
    
    for error in reversed(errors):
        salah = error.get("salah") or error.get("Kata/Frasa Salah")
        benar = error.get("benar") or error.get("Perbaikan Sesuai KBBI")
        
        if not salah or not benar:
            continue
            
        for para in doc.paragraphs:
            if salah in para.text:
                para.text = para.text.replace(salah, benar) 
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def generate_highlighted_docx(file_bytes, errors):
    doc = docx.Document(io.BytesIO(file_bytes))

    unique_salah = set(e.get("salah") or e.get("Kata/Frasa Salah") for e in errors if e.get("salah") or e.get("Kata/Frasa Salah"))
    
    for para in doc.paragraphs:
        for term in unique_salah:
            if term and term.lower() in para.text.lower():
                full_text = para.text
                para.clear()
                parts = re.split(f'({re.escape(term)})', full_text, flags=re.IGNORECASE)
                for part in parts:
                    if part:
                        run = para.add_run(part)
                        if part.lower() == term.lower():
                            run.font.color.rgb = RGBColor(255, 0, 0) 
                            run.font.bold = True 
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def create_zip_archive(revised_data, highlighted_data, original_filename):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"revisi_{original_filename}", revised_data)
        zip_file.writestr(f"highlight_{original_filename}", highlighted_data)
    return zip_buffer.getvalue()

def parse_flexible_date(date_str):
    """Mencoba parsing tanggal dari berbagai format."""
    if not date_str:
        return None
    try:
        return datetime.datetime.strptime(date_str, '%Y-%m-%dT%H:%M')
    except ValueError:
        try:
            return datetime.datetime.strptime(date_str, '%Y-%m-%d')
        except ValueError:
            raise ValueError("Format tanggal tidak valid. Gunakan format YYYY-MM-DD atau YYYY-MM-DDTHH:MM.")

def extract_paragraphs(file_bytes):
    try:
        source_stream = io.BytesIO(file_bytes)
        doc = docx.Document(source_stream)
        return [p.text for p in doc.paragraphs if p.text.strip() != ""]
    except Exception as e:
        raise ValueError(f"Gagal membaca file docx: {e}")

def extract_paragraphs_from_text(full_text):
    if not full_text:
        return []
    paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
    return paragraphs

def find_word_diff(original_para, revised_para):
    matcher = difflib.SequenceMatcher(None, original_para.split(), revised_para.split())
    diffs = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace' or tag == 'insert':
            diffs.append(" ".join(revised_para.split()[j1:j2]))
    return ", ".join(diffs) if diffs else "Perubahan Minor"

def create_comparison_docx(df):
    doc = Document()
    doc.add_heading('Hasil Perbandingan Dokumen', level=1)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df.columns):
            row_cells[i].text = str(row[col_name])
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def calculate_task_status(start_time, deadline, end_time):
    if end_time:
        return 'done'
    
    if deadline and datetime.date.today() > deadline:
        return 'overdue'
    
    return 'on_progress'

def create_recommendation_highlight_docx(file_bytes, recommendations):
    doc = docx.Document(io.BytesIO(file_bytes))
    
    misplaced_paragraphs = [rec.get("misplaced_paragraph") or rec.get("Paragraf yang Perlu Dipindah") for rec in recommendations]
    
    for para in doc.paragraphs:
        if para.text.strip() in [p.strip() for p in misplaced_paragraphs if p]:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()


def _get_word_diff_structure(original_para, revised_para):
    original_words = original_para.split()
    revised_words = revised_para.split()
    matcher = difflib.SequenceMatcher(None, original_words, revised_words)
    
    structured_output = []
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            structured_output.append({
                "text": " ".join(revised_words[j1:j2]),
                "changed": False
            })
        elif tag == 'replace' or tag == 'insert':
            structured_output.append({
                "text": " ".join(revised_words[j1:j2]),
                "changed": True
            })
    
    for item in structured_output:
        item['text'] += ' '
        
    return structured_output

def _analyze_comparison(file1, file2):
    file1_bytes = file1.read()
    file2_bytes = file2.read()
    file1.seek(0)
    file2.seek(0)

    file1_ext = file1.filename.split('.')[-1].lower()
    file2_ext = file2.filename.split('.')[-1].lower()

    pages1 = _extract_text_with_pages(file1_bytes, file1_ext)
    pages2 = _extract_text_with_pages(file2_bytes, file2_ext)
    
    original_data = extract_sentences_with_pages(pages1)
    revised_data = extract_sentences_with_pages(pages2)
    
    original_sentences = [item['sentence'] for item in original_data]
    revised_sentences = [item['sentence'] for item in revised_data]
    
    comparison_results = []
    matcher = difflib.SequenceMatcher(None, original_sentences, revised_sentences)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            len_original = i2 - i1
            len_revised = j2 - j1
            
            for i in range(min(len_original, len_revised)):
                original_sentence = original_sentences[i1 + i]
                revised_sentence = revised_sentences[j1 + i]
                revised_page = revised_data[j1 + i]['page']
                word_diff_text = find_word_diff(original_sentence, revised_sentence)
                revised_structured = _get_word_diff_structure(original_sentence, revised_sentence)
                comparison_results.append({
                    "Kalimat Awal": original_sentence,
                    "Kalimat Revisi": revised_structured,
                    "Kata yang Direvisi": word_diff_text,
                    "Halaman": f"Halaman {revised_page}"
                })
    return comparison_results

@app.route('/settings')
@login_required
def settings_page():
    return render_template('settings.html', username=current_user.username, label=current_user.label)

@app.route('/api/change_password', methods=['POST'])
@login_required
def api_change_password():
    data = request.json
    old_password = data.get('old_password')
    new_password = data.get('new_password')

    if not old_password or not new_password:
        return jsonify({"error": "Data tidak lengkap"}), 400

    if not current_user.check_password(old_password):
        return jsonify({"error": "Password lama salah"}), 400

    try:
        current_user.set_password(new_password)
        db.session.commit()
        return jsonify({"status": "success", "message": "Password berhasil diubah"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        fullname = request.form.get('fullname')
        username = request.form.get('username')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        label = request.form.get('label') 

        if not fullname or not username or not password or not confirm_password or not label:
            flash('Semua kolom wajib diisi.', 'error')
            return redirect(url_for('register'))
            
        if password != confirm_password:
            flash('Password dan Konfirmasi Password tidak cocok.', 'error')
            return redirect(url_for('register'))

        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash('Username sudah digunakan. Silakan pilih username lain.', 'error')
            return redirect(url_for('register'))
            
        try:
            new_user = User(username=username, fullname=fullname, label=label)
            new_user.set_password(password)
            db.session.add(new_user)
            db.session.commit()
            
            flash('Akun berhasil dibuat! Silakan login.', 'success')
            return redirect(url_for('login'))
            
        except Exception as e:
            db.session.rollback()
            print(f"Error Register: {e}")
            flash('Terjadi kesalahan saat membuat akun. Coba lagi nanti.', 'error')
            return redirect(url_for('register'))
            
    return render_template('register.html')

@app.route('/laporan_bulanan/<int:session_id>')
@login_required
def laporan_bulanan_page(session_id):
    """Halaman khusus Laporan Bulanan dengan logika perhitungan ganda (Prev + Curr)"""
    session = AmsTemuanSession.query.get_or_404(session_id)
    is_owner = (session.user_id == current_user.id)
    is_shared = False
    if not is_owner:
        share_check = SharedTemuanSession.query.filter_by(
            session_id=session_id, 
            shared_with_id=current_user.id
        ).first()
        if share_check:
            is_shared = True
            
    if not is_owner and not is_shared:
        return "Akses Ditolak", 403

    return render_template('laporan_bulanan.html', session=session, current_user=current_user)

@app.route('/ams_auditors')
@login_required
def ams_auditors_page():
    """Halaman Data Auditors"""
    return render_template('ams_auditors.html', 
                           username=current_user.username, 
                           label=current_user.label)

@app.route('/api/add_auditor_data', methods=['POST'])
@login_required
def api_add_auditor_data():
    """Import data auditor dari Excel via JSON"""
    req = request.get_json()
    data_list = req.get('data')
    
    if not data_list:
        return jsonify({"error": "Data kosong"}), 400

    try:
        count = 0
        for row in data_list:
            # Opsional: Bisa tambahkan logic untuk cek duplikat (update existing) jika perlu.
            # Saat ini logic-nya adalah append (tambah baru).
            new_aud = AmsAuditor(
                user_id=current_user.id,
                name=row['name'],
                department=row['department'],
                total_temuan=row['total_temuan'],
                selesai=row['selesai'],
                bjt=row['bjt'],
                outstanding=row['outstanding'],
                period=row['period']
            )
            db.session.add(new_aud)
            count += 1
        
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil menambahkan {count} data auditor."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_auditor_data', methods=['GET'])
@login_required
def api_get_auditor_data():
    """Mengambil data auditor untuk tabel, support filter periode"""
    period_filter = request.args.get('period')
    query = AmsAuditor.query.filter_by(user_id=current_user.id)
    distinct_periods = db.session.query(AmsAuditor.period).filter_by(user_id=current_user.id).distinct().all()
    all_periods = [p[0] for p in distinct_periods]
    
    if period_filter and period_filter != 'all':
        query = query.filter_by(period=period_filter)
        
    rows = query.order_by(AmsAuditor.period.desc(), AmsAuditor.name.asc()).all()
    
    result = []
    for r in rows:
        result.append({
            "name": r.name,
            "department": r.department,
            "total_temuan": r.total_temuan,
            "selesai": r.selesai,
            "bjt": r.bjt,
            "outstanding": r.outstanding,
            "period": r.period
        })
        
    return jsonify({
        "rows": result, 
        "all_periods": sorted(all_periods, reverse=True)
    })

@app.route('/api/delete_all_auditor_data', methods=['DELETE'])
@login_required
def api_delete_all_auditor_data():
    """Menghapus semua data auditor milik user ini"""
    try:
        AmsAuditor.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500
    
@app.route('/api/get_auditor_performance', methods=['GET'])
@login_required
def api_get_auditor_performance():
    """
    Menghitung performa auditor:
    Delta Selesai = (Selesai Bulan Ini) - (Selesai Bulan Sebelumnya)
    """
    try:
        periods_query = db.session.query(AmsAuditor.period).filter_by(user_id=current_user.id).distinct().all()
        periods = sorted([p[0] for p in periods_query], reverse=True)
        
        if not periods:
            return jsonify([]) 
            
        current_period = periods[0] 
        prev_period = periods[1] if len(periods) > 1 else None 

        current_data = AmsAuditor.query.filter_by(user_id=current_user.id, period=current_period).all()
        curr_map = {d.name: d.selesai for d in current_data}
        
        prev_map = {}
        if prev_period:
            prev_data = AmsAuditor.query.filter_by(user_id=current_user.id, period=prev_period).all()
            prev_map = {d.name: d.selesai for d in prev_data}
            
        result = []
        for name, curr_val in curr_map.items():
            prev_val = prev_map.get(name, 0) 
            delta = curr_val - prev_val
            
            result.append({
                "name": name,
                "delta": delta,
                "current": curr_val,
                "prev": prev_val
            })
            
        return jsonify(result)

    except Exception as e:
        print(f"Error calculating auditor performance: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/library')
@login_required
def library_page():
    return render_template('library.html', current_user=current_user)

@app.route('/api/library/list', methods=['GET'])
@login_required
def api_library_list():
    try:
        files = LibraryFile.query.filter_by(user_id=current_user.id).order_by(LibraryFile.upload_date.desc()).all()
        data = []
        for f in files:
            data.append({
                "id": f.id,
                "title": f.title,
                "cluster": f.cluster,
                "shelf_id": f.shelf_id,
                "summary": f.summary,
                "category": f.category,
                "size": f.file_size,
                "type": f.file_type,
                "url": url_for('api_library_view', file_id=f.id, filename=f.filename)
            })
        return jsonify(data), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/library/upload', methods=['POST'])
@login_required
def api_library_upload():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file yang diunggah"}), 400
    
    file = request.files['file']
    title = request.form.get('title')
    category = request.form.get('category')
    summary = request.form.get('summary')
    cluster = request.form.get('cluster')
    shelf_id = request.form.get('shelf_id')

    if file.filename == '':
        return jsonify({"error": "Nama file kosong"}), 400

    try:
        user_folder = get_user_root_folder()
        library_folder = os.path.join(user_folder, 'Library_Files (Jangan Dihapus)')
        os.makedirs(library_folder, exist_ok=True)
        filename = secure_filename(file.filename)
        save_path = os.path.join(library_folder, filename)
        
        if os.path.exists(save_path):
            timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"{timestamp}_{filename}"
            save_path = os.path.join(library_folder, filename)
            
        file.save(save_path)
        file_size_bytes = os.path.getsize(save_path)
        
        if file_size_bytes < 1024:
            size_str = f"{file_size_bytes} B"
        elif file_size_bytes < 1024 * 1024:
            size_str = f"{file_size_bytes / 1024:.1f} KB"
        else:
            size_str = f"{file_size_bytes / (1024 * 1024):.1f} MB"

        file_ext = filename.split('.')[-1].lower()

        new_file = LibraryFile(
            user_id=current_user.id,
            title=title,
            category=category,
            summary=summary,
            cluster=cluster,
            filename=filename,
            file_path=save_path,
            file_size=size_str,
            file_type=file_ext,
            shelf_id=shelf_id
        )
        db.session.add(new_file)
        db.session.commit()

        return jsonify({"status": "success", "message": "File berhasil diupload"}), 201

    except Exception as e:
        db.session.rollback()
        print(f"Error Upload Library: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/library/delete/<int:file_id>', methods=['DELETE'])
@login_required
def api_library_delete(file_id):
    file_record = LibraryFile.query.get(file_id)
    if not file_record or file_record.user_id != current_user.id:
        return jsonify({"error": "File tidak ditemukan"}), 404
    
    try:
        if os.path.exists(file_record.file_path):
            os.remove(file_record.file_path)

        db.session.delete(file_record)
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/library/view/<int:file_id>/<path:filename>')
@login_required
def api_library_view(file_id, filename):
    file_record = LibraryFile.query.get(file_id)
    
    if not file_record or file_record.user_id != current_user.id:
        return "File not found or access denied", 404
    return send_file(
        file_record.file_path, 
        as_attachment=False, 
        download_name=file_record.filename
    )

@app.route('/', methods=['GET', 'POST']) 
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard')) 
    
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash('Username atau password salah. Harap periksa kembali kesesuaian username dan password Anda.') 
            
    return render_template('login.html')

@app.route('/api/save_gsheet_url', methods=['POST'])
@login_required
def api_save_gsheet_url():
    data = request.json
    session_id = data.get('session_id')
    url = data.get('url')
    
    session = AmsTemuanSession.query.get(session_id)
    if not session or session.user_id != current_user.id:
        return jsonify({"error": "Sesi tidak ditemukan / Akses ditolak"}), 403
        
    session.google_sheet_url = url
    db.session.commit()
    return jsonify({"status": "success", "message": "Link Google Sheet tersimpan."})

@app.route('/logout')
@login_required 
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/mailbox') 
@login_required
def mailbox_page():
    """Halaman utama untuk fitur Kotak Surat (Email)"""
    return render_template('mailbox.html',current_user=current_user, username=current_user.username,
                           label=current_user.label)

@app.route('/dashboard') 
@login_required 
def dashboard(): 
    """
    REVISI: Mengambil data folder (objek) dari get_user_folders.
    """
    folders = []
    try:
        folders = get_user_folders() 
    except Exception as e:
        print(f"Error memuat folder untuk dashboard: {e}")
        flash("Gagal memuat struktur folder Anda.")
    return render_template('index.html', 
                            username=current_user.username, 
                            label=current_user.label,
                            folders=folders,
                            current_user=current_user,
                            current_user_id=current_user.id)

@app.route('/api/create_folder', methods=['POST'])
@login_required 
def api_create_folder():
    """Membuat folder baru di sistem file."""
    data = request.json
    name = data.get('name')

    if not name:
        return jsonify({"error": "Nama folder tidak boleh kosong."}), 400

    try:
        clean_name = create_user_folder(name)
        return jsonify({
            "status": "success",
            "folder_name": clean_name
        }), 201
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        print(f"Kesalahan Server saat membuat folder: {e}")
        return jsonify({"error": "Gagal membuat folder di server."}), 500

@app.route('/api/list_folders', methods=['GET'])
@login_required 
def api_list_folders():
    """Mendaftarkan semua folder yang dimiliki pengguna."""
    print(f"Mencoba mengambil folder untuk user: {current_user.id}") 
    try:
        folders = get_user_folders() 
        print(f"Ditemukan folder: {folders}") 
        return jsonify(folders)
    except Exception as e:
        print(f"Kesalahan Server saat mengambil list folder: {e}")
        return jsonify({"error": "Gagal memuat folder: " + str(e)}), 500

@app.route('/api/delete_folder', methods=['POST'])
@login_required
def api_delete_folder():
    """Menghapus folder dan seluruh isinya."""
    data = request.json
    folder_name = data.get('folder_name')

    if not folder_name or '..' in folder_name or '~' in folder_name:
        return jsonify({"error": "Nama folder tidak valid."}), 400
    
    try:
        user_root = get_user_root_folder()
        if not user_root:
            return jsonify({"error": "User tidak ditemukan."}), 401
            
        folder_path = os.path.join(user_root, folder_name)
        if not os.path.isdir(folder_path) or not os.path.abspath(folder_path).startswith(os.path.abspath(user_root)):
             return jsonify({"error": "Folder tidak ditemukan atau akses ditolak."}), 404
        shutil.rmtree(folder_path)
        SharedFolder.query.filter_by(owner_id=current_user.id, folder_name=folder_name).delete()
        db.session.commit()
        
        return jsonify({"status": "success", "message": f"Folder '{folder_name}' berhasil dihapus."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus folder: {e}")
        return jsonify({"error": f"Gagal menghapus folder: {e}"}), 500


@app.route('/api/save_results', methods=['POST'])
@login_required 
def api_save_results():
    """Menyimpan hasil analisis ke file JSON di dalam folder."""
    data = request.json
    folder_name = data.get('folder_name')
    feature_type = data.get('feature_type')
    results_data = data.get('results_data')
    original_filename = data.get('original_filename', 'untitled_analysis')
    owner_id = data.get('owner_id', current_user.id)
    actions_data = data.get('actions_data', {})

    if not folder_name or not feature_type or not results_data:
        return jsonify({"error": "Data folder, fitur, atau hasil kosong."}), 400

    try:
        target_user_id_str = str(owner_id)
        if not target_user_id_str.isalnum():
             return jsonify({"error": "Owner ID tidak valid."}), 400
        
        user_root = os.path.join(app.config['UPLOAD_FOLDER'], target_user_id_str)
        if '..' in folder_name:
             return jsonify({"error": "Nama folder tidak valid."}), 400
        folder_path = os.path.join(user_root, folder_name)
        is_owner = (str(current_user.id) == target_user_id_str)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak untuk menyimpan ke folder ini."}), 403
        
        if not os.path.isdir(folder_path):
            return jsonify({"error": "Folder tidak ditemukan."}), 404
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_orig_name = re.sub(r'[^\w\s-]', '', original_filename.split('.')[0]).strip()
        clean_orig_name = re.sub(r'[-\s]+', '_', clean_orig_name)
        
        save_filename = f"{timestamp}_{feature_type}_{clean_orig_name}.json"
        save_path = os.path.join(folder_path, save_filename)
        with open(save_path, 'w', encoding='utf-8') as f:
            json.dump(results_data, f, ensure_ascii=False, indent=2)
        for row_id_str, action_data in actions_data.items():
            try:
                row_id = int(row_id_str) 
                action = RowAction.query.filter_by(
                    owner_id=owner_id,
                    folder_name=folder_name,
                    filename=save_filename, 
                    row_id=row_id
                ).first()

                if action:
                    action.is_ganti = action_data.get('is_ganti', False)
                    action.pic_user_id = action_data.get('pic_user_id')
                else:
                    action = RowAction(
                        owner_id=owner_id,
                        folder_name=folder_name,
                        filename=save_filename,
                        row_id=row_id,
                        is_ganti=action_data.get('is_ganti', False),
                        pic_user_id=action_data.get('pic_user_id')
                    )
                    db.session.add(action)
            except Exception as e:
                print(f"Gagal memproses aksi untuk baris {row_id_str}: {e}")
                raise e 

        db.session.commit()

        return jsonify({
            "status": "success",
            "message": f"Hasil analisis dan status aksi tersimpan di folder {folder_name}."
        }), 201
    except Exception as e:
        db.session.rollback()
        print(f"Kesalahan Server saat menyimpan hasil: {e}")
        return jsonify({"error": "Gagal menyimpan hasil: " + str(e)}), 500

@app.route('/api/get_all_users', methods=['GET'])
@login_required
def api_get_all_users():
    try:
        users = User.query.all()
        user_list = [{"id": user.id, "username": user.username, "fullname": user.fullname, "label": user.label} for user in users]
        return jsonify(user_list)
    except Exception as e:
        print(f"Error get all users: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/share_folder', methods=['POST'])
@login_required
def api_share_folder():
    """Berbagi folder dengan user lain."""
    data = request.json
    folder_name = data.get('folder_name')
    share_with_user_ids = data.get('share_with_user_ids')

    if not folder_name or not share_with_user_ids:
        return jsonify({"error": "Data tidak lengkap (folder atau user)."}), 400
    
    if not isinstance(share_with_user_ids, list):
         return jsonify({"error": "Format user ID harus berupa list."}), 400

    success_names = []  
    skipped_count = 0
    errors = []

    for user_id_str in share_with_user_ids: 
        try:
            user_id = int(user_id_str) 
        except ValueError:
            errors.append(f"User ID tidak valid: {user_id_str}")
            continue
    
        user_to_share = User.query.get(user_id)
        if not user_to_share:
            errors.append(f"Pengguna dengan ID {user_id} tidak ditemukan.")
            continue

        existing_share = SharedFolder.query.filter_by(
            owner_id=current_user.id,
            folder_name=folder_name,
            shared_with_id=user_id
        ).first()

        if existing_share:
            skipped_count += 1
            continue

        try:
            new_share = SharedFolder(
                owner_id=current_user.id,
                folder_name=folder_name,
                shared_with_id=user_id,
                folder_type='general' 
            )
            db.session.add(new_share)
            success_names.append(user_to_share.username) 
        
        except Exception as e:
            db.session.rollback() 
            errors.append(f"Gagal share ke {user_to_share.username}: {e}")
            continue 

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Gagal menyimpan ke database: {e}"}), 500

    return jsonify({
        "status": "success",
        "message": f"Berhasil di-share ke {len(success_names)} user.",
        "success_names": success_names,  
        "skipped_count": skipped_count,
        "errors": errors
    }), 201
    

@app.route('/api/folder_history/<int:owner_id>/<folder_name>', methods=['GET'])
@login_required
def api_folder_history(owner_id, folder_name):
    """
    PERBAIKAN UNTUK MASALAH 1: Mengambil riwayat file dari folder.
    Kita butuh owner_id untuk folder yang di-share.
    """
    if not folder_name or '..' in folder_name:
        return jsonify({"error": "Nama folder tidak valid."}), 400
    
    try:
        target_user_id_str = str(owner_id)
        if not target_user_id_str.isalnum():
             return jsonify({"error": "Owner ID tidak valid."}), 400
        
        folder_root = os.path.join(app.config['UPLOAD_FOLDER'], target_user_id_str)
        folder_path = os.path.join(folder_root, folder_name)

        is_owner = (str(current_user.id) == target_user_id_str)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak."}), 403
            
        if not os.path.isdir(folder_path):
            return jsonify({"error": "Folder tidak ditemukan."}), 404

        results = []
        for filename in os.listdir(folder_path):
            if filename.endswith('.json'):
                try:
                    parts = filename.replace('.json', '').split('_', 3)
                    timestamp_str = f"{parts[0]}_{parts[1]}"
                    timestamp = datetime.datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
                    feature_type = parts[2]
                    original_name = parts[3] if len(parts) > 3 else "N/A"
                    
                    results.append({
                        "filename": filename,
                        "feature_type": feature_type,
                        "timestamp": timestamp.strftime("%d %b %Y, %H:%M"),
                        "original_name": original_name.replace("_", " ")
                    })
                except Exception as e:
                    print(f"Gagal parse nama file {filename}: {e}")
                    results.append({
                        "filename": filename,
                        "feature_type": "N/A",
                        "timestamp": "N/A",
                        "original_name": "N/A"
                    })
        
        results.sort(key=lambda x: x.get("timestamp"), reverse=True)
        
        return jsonify(results)

    except Exception as e:
        print(f"Error saat mengambil riwayat folder: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/mailbox/view/<int:message_id>')
@login_required
def view_message_detail(message_id):
    message = Message.query.get_or_404(message_id)
    if message.sender_id != current_user.id and message.recipient_id != current_user.id:
        flash("Akses ditolak ke pesan ini.")
        return redirect(url_for('mailbox_page'))
    if message.recipient_id == current_user.id and not message.is_read:
        message.is_read = True
        db.session.commit()
    return render_template('view_message.html', message=message, current_user=current_user)

@app.route('/api/get_result_file', methods=['POST'])
@login_required
def api_get_result_file():
    data = request.json
    folder_name = data.get('folder_name')
    filename = data.get('filename')
    
    if not folder_name or not filename:
        return jsonify({"error": "Data tidak lengkap."}), 400
    
    if '..' in folder_name or '..' in filename:
        return jsonify({"error": "Nama file/folder tidak valid."}), 400

    try:
        try:
            owner_id = int(data.get('owner_id'))
        except (TypeError, ValueError):
            return jsonify({"error": "Owner ID tidak valid."}), 400

        print(f"[DEBUG GET] Fetching: Owner={owner_id}, Folder={folder_name}, File={filename}")

        target_user_id_str = str(owner_id)
        
        is_owner = (current_user.id == owner_id)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak."}), 403
            
        folder_root = os.path.join(app.config['UPLOAD_FOLDER'], target_user_id_str)
        file_path = os.path.join(folder_root, folder_name, filename)

        if not os.path.isfile(file_path):
            return jsonify({"error": "File tidak ditemukan."}), 404

        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
            
        row_actions = RowAction.query.filter_by(
            owner_id=owner_id,
            folder_name=folder_name,
            filename=filename
        ).all()

        print(f"[DEBUG GET] Found {len(row_actions)} saved actions in DB.")

        actions_data = {
            action.row_id: {
                'is_ganti': action.is_ganti, 
                'pic_user_id': action.pic_user_id
            } for action in row_actions
        }

        response = jsonify({
            "status": "success",
            "data": json_data,
            "actions": actions_data
        })
    
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        
        return response

    except Exception as e:
        print(f"[ERROR GET] {e}")
        return jsonify({"error": f"Server Error: {str(e)}"}), 500

@app.route('/api/proofread/analyze', methods=['POST'])
@login_required 
def api_proofread_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file yang diunggah"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nama file kosong"}), 400

    try:
        file_bytes = file.read()
        file.seek(0)
        file_extension = file.filename.split('.')[-1].lower()
        try:
            pages_content = _extract_text_with_pages(file_bytes, file_extension)
        except ValueError as ve:
            return jsonify({"error": str(ve)}), 400

        all_errors = []
        for page in pages_content:
            page_text = page['teks']
            page_num = page['halaman']
            
            page_errors = proofread_with_gemini(page_text)
            for error in page_errors:
                formatted_error = {
                    "Kata/Frasa Salah": error.get('salah', ''),
                    "Perbaikan Sesuai KBBI": error.get('benar', ''),
                    "Pada Kalimat": error.get('kalimat', ''),
                    "Ditemukan di Halaman": f"Halaman {page_num}",
                    "apakah_ganti": False, 
                    "salah": error.get('salah', ''),
                    "benar": error.get('benar', '')
                }
                all_errors.append(formatted_error)

        try:
            new_log = AnalysisLog(
                user_id=current_user.id,
                filename=file.filename,
                feature_type='proofreading',
                status='done',
                end_time=datetime.datetime.utcnow()
            )
            db.session.add(new_log)
            db.session.commit()
        except Exception as e:
            print(f"Gagal menyimpan log: {e}")

        return jsonify(all_errors)

    except Exception as e:
        print(f"[ERROR PROOFREAD] {e}") 
        return jsonify({"error": f"Terjadi kesalahan server: {str(e)}"}), 500

@app.route('/api/delete_result', methods=['POST'])
@login_required
def api_delete_result():
    
    """Menghapus satu file hasil analisis (JSON) dari folder."""
    data = request.json
    folder_name = data.get('folder_name')
    filename = data.get('filename')
    owner_id = data.get('owner_id') 

    if not folder_name or not filename or not owner_id:
        return jsonify({"error": "Data tidak lengkap."}), 400
    
    if '..' in folder_name or '..' in filename:
        return jsonify({"error": "Nama file/folder tidak valid."}), 400

    try:
        if str(current_user.id) != str(owner_id):
             return jsonify({"error": "Akses ditolak. Hanya pemilik folder yang bisa menghapus file."}), 403

        user_root = get_user_root_folder() 
        file_path = os.path.join(user_root, folder_name, filename)

        if not os.path.isfile(file_path) or not os.path.abspath(file_path).startswith(os.path.abspath(user_root)):
            return jsonify({"error": "File tidak ditemukan atau akses ditolak."}), 404
        
        os.remove(file_path)
        
        return jsonify({"status": "success", "message": f"File '{filename}' berhasil dihapus."}), 200

    except Exception as e:
        print(f"Error saat menghapus file: {e}")
        return jsonify({"error": f"Gagal menghapus file: {e}"}), 500

def _generate_proofread_files(file, file_bytes):
    """Helper internal untuk download, menjalankan analisis lagi."""
    file_extension = file.filename.split('.')[-1].lower()
    document_pages = _extract_text_with_pages(file_bytes, file_extension)
    
    all_errors = []
    for page in document_pages:
        found_errors_on_page = proofread_with_gemini(page['teks'])
        all_errors.extend(found_errors_on_page) 

    revised_data = generate_revised_docx(file_bytes, all_errors)
    highlighted_data = generate_highlighted_docx(file_bytes, all_errors)
    
    return revised_data, highlighted_data, file.filename

@app.route('/api/proofread/download/revised', methods=['POST'])
@login_required 
def api_proofread_download_revised():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() 
    file.seek(0) 
    
    try:
        revised_data, _, filename = _generate_proofread_files(file, file_bytes)
        
        return send_file(
            io.BytesIO(revised_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"revisi_{filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/proofread/download/highlighted', methods=['POST'])
@login_required 
def api_proofread_download_highlighted():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() 
    file.seek(0)
    
    try:
        _, highlighted_data, filename = _generate_proofread_files(file, file_bytes)
        
        return send_file(
            io.BytesIO(highlighted_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"highlight_{filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/proofread/download/zip', methods=['POST'])
@login_required 
def api_proofread_download_zip():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() 
    file.seek(0)
    
    try:
        revised_data, highlighted_data, filename = _generate_proofread_files(file, file_bytes)
        zip_data = create_zip_archive(revised_data, highlighted_data, filename)
        
        return send_file(
            io.BytesIO(zip_data),
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"hasil_proofread_{filename.split('.')[0]}.zip"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/analyze_advanced', methods=['POST'])
@login_required 
def api_compare_analyze_advanced():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        full_text1 = _get_full_text_from_file(file1)
        full_text2 = _get_full_text_from_file(file2)
        comparison_results_from_ai = analyze_document_by_section(full_text1, full_text2)
        
        final_results = []
        for item in comparison_results_from_ai:
            sub_bab_asli = item.get("sub_bab_asal", "N/A")
            
            try:
                nama_sub_bab = sub_bab_asli.split(':', 1)[1].strip()
            except (IndexError, AttributeError):
                nama_sub_bab = sub_bab_asli
            
            final_results.append({
                "Sub-bab Referensi pada Dokumen asli": item.get("sub_bab_referensi", "-"),
                "Sub-bab Asal (Pada dokumen yang dibanding)": nama_sub_bab,
                "Kalimat Menyimpang (Dokumen yang dibanding)": item.get("kalimat_menyimpang", "N/A"),
                "Alasan": item.get("alasan", "N/A")
            })
        
        return jsonify(final_results)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/analyze', methods=['POST'])
@login_required 
def api_compare_analyze():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        results = _analyze_comparison(file1, file2)
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/download', methods=['POST'])
@login_required 
def api_compare_download():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        original_paras = extract_paragraphs(file1.read())
        revised_paras = extract_paragraphs(file2.read())
        comparison_results = []
        matcher = difflib.SequenceMatcher(None, original_paras, revised_paras)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                len_original = i2 - i1
                len_revised = j2 - j1
                for i in range(min(len_original, len_revised)):
                    original_para = original_paras[i1 + i]
                    revised_para = revised_paras[j1 + i]
                    word_diff = find_word_diff(original_para, revised_para)
                    comparison_results.append({
                        "Kalimat Awal": original_para,
                        "Kalimat Revisi": revised_para,
                        "Kata yang Direvisi": word_diff,
                    })
        
        if not comparison_results:
             return jsonify({"error": "Tidak ada perbedaan untuk diunduh"}), 400

        df_comparison = pd.DataFrame(comparison_results)
        docx_data = create_comparison_docx(df_comparison)
        
        return send_file(
            io.BytesIO(docx_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"perbandingan_{file1.filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/coherence/analyze', methods=['POST'])
@login_required 
def api_coherence_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        full_text = _get_full_text_from_file(file)
        issues_from_gemini = analyze_document_coherence(full_text)
        
        processed_issues = []
        for issue in issues_from_gemini:
            asli_text = issue['asli']
            saran_text = issue['saran']
            
            saran_structured = _get_word_diff_structure(asli_text, saran_text)
            
            processed_issues.append({
                "topik": issue['topik'],
                "asli": asli_text,
                "saran": saran_structured,
                "catatan": issue['catatan']
            })
            
        return jsonify(processed_issues)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _analyze_restructure(file):
    full_text = _get_full_text_from_file(file)
    recommendations = get_structural_recommendations(full_text)
    processed_results = []
    for rec in recommendations:
        processed_results.append({
            "Paragraf yang Perlu Dipindah": rec.get("misplaced_paragraph"),
            "Lokasi Asli": rec.get("original_section"),
            "Saran Lokasi Baru": rec.get("recommended_section")
        })
    return processed_results

@app.route('/api/restructure/analyze', methods=['POST'])
@login_required 
def api_restructure_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        results = _analyze_restructure(file)
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/restructure/download', methods=['POST'])
@login_required 
def api_restructure_download():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() 
    file.seek(0) 
    
    try:
        full_text = "\n".join([p['teks'] for p in _extract_text_with_pages(file_bytes, file.filename.split('.')[-1].lower())])
        recommendations = get_structural_recommendations(full_text)
        processed_results = []
        for rec in recommendations:
            processed_results.append({
                "Paragraf yang Perlu Dipindah": rec.get("misplaced_paragraph") 
            })

        if not processed_results or "Error:" in processed_results[0].get("Paragraf yang Perlu Dipindah", ""):
             return jsonify({"error": "Tidak ada rekomendasi valid untuk diunduh"}), 400

        highlighted_data = create_recommendation_highlight_docx(file_bytes, processed_results)
        
        return send_file(
            io.BytesIO(highlighted_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"highlight_rekomendasi_{file.filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/review/analyze', methods=['POST'])
@login_required 
def api_review_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file yang diunggah"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nama file kosong"}), 400

    try:
        full_text = _get_full_text_from_file(file)
        results = review_document_comprehensive(full_text)
        try:
            new_log = AnalysisLog(
                user_id=current_user.id,
                filename=file.filename,
                feature_type='review_dokumen',
                status='done',
                end_time=datetime.datetime.utcnow()
            )
            db.session.add(new_log)
            db.session.commit()
        except Exception as e:
            print(f"Gagal simpan log: {e}")

        return jsonify(results)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_comments', methods=['POST'])
@login_required
def api_get_comments():
    data = request.json
    folder_name = data.get('folder_name')
    filename = data.get('filename')
    if not folder_name or not filename:
        return jsonify({"error": "Data file tidak lengkap."}), 400
    all_comments = Comment.query.filter_by(
        folder_name=folder_name,
        filename=filename
    ).order_by(Comment.timestamp.asc()).all()
    comments_by_id = {c.id: c for c in all_comments}
    top_level_comments = [c for c in all_comments if c.parent_id is None]

    def build_comment_tree(comment):
        comment_dict = {
            'id': comment.id,
            'row_id': comment.row_id,
            'username': comment.username,
            'text': comment.text,
            'timestamp': comment.timestamp.isoformat(),
            'replies': []
        }
        for reply in all_comments:
            if reply.parent_id == comment.id:
                comment_dict['replies'].append(build_comment_tree(reply))
        return comment_dict

    nested_comments = [build_comment_tree(c) for c in top_level_comments]

    return jsonify(nested_comments), 200

@app.route('/api/edit_auditee_data/<int:auditee_id>', methods=['POST'])
@login_required
def api_edit_auditee_data(auditee_id):
    """Edit data auditee."""
    data = request.get_json()
    auditee = AmsMonitoring.query.get(auditee_id)
    
    if not auditee or auditee.user_id != current_user.id:
        return jsonify({"error": "Data tidak ditemukan atau akses ditolak."}), 404

    try:
        auditee.auditee = data.get('auditee')
        auditee.tahun_audit = int(data.get('tahun_audit'))
        auditee.total_rekomendasi = int(data.get('total_rekomendasi', 0))
        auditee.selesai = int(data.get('selesai', 0))
        auditee.tidak_selesai = int(data.get('tidak_selesai', 0))
        auditee.todo = int(data.get('todo', 0))
        auditee.belum_sesuai = int(data.get('belum_sesuai', 0))
        auditee.belum_tl = int(data.get('belum_tl', 0))
        auditee.tdd = int(data.get('tdd', 0))
        
        db.session.commit()
        return jsonify({"status": "success", "message": "Data auditee berhasil diperbarui."}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error update auditee: {e}")
        return jsonify({"error": "Gagal memperbarui data auditee."}), 500

@app.route('/api/add_monitoring_session', methods=['POST'])
@login_required
def api_add_monitoring_session():
    """Mempersiapkan sesi baru (Validasi nama & tipe)."""
    data = request.get_json()
    session_name = data.get('session_name')
    monitoring_type = data.get('monitoring_type', 'standard') 

    if not session_name:
        return jsonify({"error": "Nama sesi tidak boleh kosong."}), 400
    existing = AmsMonitoring.query.filter_by(user_id=current_user.id, session_name=session_name).first()
    if existing:
        return jsonify({"error": "Nama sesi sudah digunakan. Mohon pilih nama lain."}), 400

    return jsonify({
        "status": "success",
        "message": f"Sesi '{session_name}' siap digunakan.",
        "session_name": session_name,
        "monitoring_type": monitoring_type
    }), 201

@app.route('/monitoring_ams/view/<path:session_name>')
@login_required
def monitoring_view_page(session_name):
    """Halaman khusus Full Screen untuk satu sesi monitoring"""
    owner_id = request.args.get('owner_id')
    return render_template('monitoring_view.html', 
                           session_name=session_name, 
                           owner_id=owner_id,
                           username=current_user.username)

@app.route('/api/share_ams_session', methods=['POST'])
@login_required
def api_share_ams_session():
    """Membagikan sesi AMS ke user lain."""
    data = request.json
    session_name = data.get('session_name')
    target_user_ids = data.get('user_ids', [])

    if not session_name or not target_user_ids:
        return jsonify({"error": "Data tidak lengkap."}), 400

    is_owner = AmsMonitoring.query.filter_by(user_id=current_user.id, session_name=session_name).first()
    if not is_owner:
        return jsonify({"error": "Anda bukan pemilik sesi ini."}), 403

    success_count = 0
    for uid in target_user_ids:
        try:
            exists = SharedMonitoring.query.filter_by(
                owner_id=current_user.id,
                shared_with_id=uid,
                session_name=session_name
            ).first()

            if not exists:
                new_share = SharedMonitoring(
                    owner_id=current_user.id,
                    shared_with_id=uid,
                    session_name=session_name
                )
                db.session.add(new_share)
                success_count += 1
        except Exception:
            continue

    try:
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil dibagikan ke {success_count} pengguna."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload_ams_image', methods=['POST'])
@login_required
def api_upload_ams_image():
    # 1. Validasi Input Dasar
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file gambar."}), 400
    
    file = request.files['file']
    session_name = request.form.get('session_name', '')
    
    if not file or not session_name:
        return jsonify({"error": "File atau nama sesi hilang."}), 400

    print(f"\n[DEBUG] Memulai Upload Gambar untuk Sesi: {session_name}")

    try:
        # 2. Proses Gambar ke Gemini
        img = PIL.Image.open(file)
        
        # Tentukan Mode Prompt
        is_bpk_mode = False
        if 'bpk' in session_name.lower() or 'assessment' in session_name.lower():
            is_bpk_mode = True

        prompt = """
        Analisis gambar tabel audit ini. Ekstrak data baris per baris menjadi JSON Array.
        
        Aturan Penting:
        1. HANYA kembalikan JSON Array murni [ ... ]. Jangan ada markdown (```json), jangan ada teks pengantar.
        2. Pastikan setiap objek memiliki key berikut (gunakan lowercase):
           - "auditee": Nama unit atau auditee.
           - "tahun_audit": Tahun (angka).
           - "total_rekomendasi": Total rekomendasi (angka).
           - "selesai": Jumlah selesai (angka).
           - "tidak_selesai": (Khusus standar) Belum jatuh tempo (angka).
           - "todo": (Khusus standar) Outstanding (angka).
           - "belum_sesuai": (Khusus BPK) Belum sesuai (angka).
           - "belum_tl": (Khusus BPK) Belum tindak lanjut (angka).
           - "tdd": (Khusus BPK) Tidak dapat ditindaklanjuti (angka).
        3. Jika kolom berupa persentase (%), AMBIL ANGKA JUMLAHNYA SAJA yang ada di sebelahnya.
        """

        model = genai.GenerativeModel('gemini-2.5-flash') # Gunakan Flash agar cepat, atau Pro jika butuh akurasi tinggi
        track_gemini_usage()
        
        print("[DEBUG] Mengirim ke Gemini...")
        response = model.generate_content([prompt, img])
        raw_text = response.text
        
        print(f"[DEBUG] Raw Response dari AI:\n{raw_text}\n-------------------")

        # 3. Pembersihan JSON (Cleaning)
        # Hapus markdown ```json dan ```
        clean_json = raw_text.replace('```json', '').replace('```', '').strip()
        
        # Cari kurung siku pertama '[' dan terakhir ']' untuk isolasi array
        start_idx = clean_json.find('[')
        end_idx = clean_json.rfind(']')
        
        if start_idx == -1 or end_idx == -1:
             raise ValueError("AI tidak mengembalikan JSON Array yang valid (kurung siku [] tidak ditemukan).")
        
        clean_json = clean_json[start_idx:end_idx+1]

        try:
            extracted_data = json.loads(clean_json)
        except json.JSONDecodeError as e:
            print(f"[ERROR] JSON Parsing Gagal: {e}")
            raise ValueError("Gagal membaca format data dari AI.")

        if not extracted_data:
            raise ValueError("Hasil ekstraksi kosong (array []). Coba gambar yang lebih jelas.")

        # 4. Helper Function: Normalisasi Key & Angka
        # Fungsi ini mencari value meskipun AI salah ketik huruf besar/kecil
        def get_val(item, keys_list):
            # Buat dict lowercase
            item_lower = {k.lower(): v for k, v in item.items()}
            for k in keys_list:
                if k.lower() in item_lower:
                    return item_lower[k.lower()]
            return 0 # Default jika tidak ketemu

        def clean_int(val):
            if not val: return 0
            # Ubah ke string, hapus simbol non-angka (kecuali minus), lalu ambil angka pertama
            s = str(val).replace('%', '').replace(',', '').replace('.', '')
            import re
            match = re.search(r'\d+', s)
            return int(match.group()) if match else 0

        # 5. Looping Simpan ke DB
        count = 0
        for item in extracted_data:
            # Ambil Data dengan aman (Key Insensitive)
            auditee_val = get_val(item, ['auditee', 'nama_auditee', 'unit']) or 'Unknown'
            tahun_val = clean_int(get_val(item, ['tahun', 'tahun_audit']))
            if tahun_val < 2000: tahun_val = datetime.date.today().year # Fallback tahun

            total_rek = clean_int(get_val(item, ['total', 'total_rekomendasi', 'jumlah']))
            
            selesai_val = clean_int(get_val(item, ['selesai', 'done']))
            
            # Variabel Khusus BPK / Standard
            tidak_selesai_val = clean_int(get_val(item, ['tidak_selesai', 'bjt', 'belum_jatuh_tempo']))
            todo_val = clean_int(get_val(item, ['todo', 'outstanding', 'out']))
            
            belum_sesuai_val = clean_int(get_val(item, ['belum_sesuai', 'bs', 'belum_sesuai_rekomendasi']))
            belum_tl_val = clean_int(get_val(item, ['belum_tl', 'btl', 'belum_tindak_lanjut']))
            tdd_val = clean_int(get_val(item, ['tdd', 'tidak_dapat_ditindaklanjuti']))

            detected_type = 'bpk' if is_bpk_mode else 'standard'

            new_entry = AmsMonitoring(
                user_id=current_user.id,
                session_name=session_name,
                monitoring_type=detected_type,
                auditee=str(auditee_val),
                tahun_audit=tahun_val,
                total_rekomendasi=total_rek,
                selesai=selesai_val,
                tidak_selesai=tidak_selesai_val,
                todo=todo_val,
                belum_sesuai=belum_sesuai_val,
                belum_tl=belum_tl_val,
                tdd=tdd_val
            )
            db.session.add(new_entry)
            count += 1
            print(f"[DEBUG] Menyiapkan data: {auditee_val} - Total: {total_rek}")

        # 6. Commit Database
        db.session.commit()
        print(f"[SUCCESS] Berhasil menyimpan {count} data ke database.")
        
        return jsonify({
            "status": "success", 
            "message": f"Berhasil mengekstrak {count} data. Silakan cek tabel."
        }), 200

    except Exception as e:
        db.session.rollback()
        print(f"[CRITICAL ERROR] {str(e)}")
        import traceback
        traceback.print_exc() # Print full error log di terminal
        return jsonify({"error": f"Gagal: {str(e)}"}), 500

@app.route('/api/get_monitoring_sessions', methods=['GET'])
@login_required
def api_get_monitoring_sessions():
    """Mengambil daftar sesi beserta tipenya (Versi Stabil)."""
    try:
        own_sessions = db.session.query(
            AmsMonitoring.session_name, 
            AmsMonitoring.monitoring_type
        ).filter_by(user_id=current_user.id).distinct().all()
        
        result = []

        for s in own_sessions:
            result.append({
                "name": s[0], 
                "type": s[1], 
                "owner": None, 
                "is_shared": False
            })

        shared_sessions = SharedMonitoring.query.filter_by(shared_with_id=current_user.id).all()
        for share in shared_sessions:
            owner = User.query.get(share.owner_id)
            owner_name = owner.username if owner else "Unknown"

            sample = AmsMonitoring.query.filter_by(
                user_id=share.owner_id, 
                session_name=share.session_name
            ).first()
            
            sType = sample.monitoring_type if sample else 'standard'

            result.append({
                "name": share.session_name,
                "type": sType,
                "owner": owner_name,
                "is_shared": True,
                "owner_id": share.owner_id
            })

        return jsonify(result), 200
    except Exception as e:
        print(f"[ERROR API SESSION] {e}") 
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_monitoring_data/<session_name>', methods=['GET'])
@login_required
def api_get_monitoring_data(session_name):
    try:
        owner_id_param = request.args.get('owner_id')
        target_user_id = current_user.id

        if owner_id_param and owner_id_param != 'null':
             has_access = SharedMonitoring.query.filter_by(
                 owner_id=owner_id_param, shared_with_id=current_user.id, session_name=session_name
             ).first()
             if not has_access: return jsonify({"error": "Akses ditolak."}), 403
             target_user_id = owner_id_param

        auditees = AmsMonitoring.query.filter_by(user_id=target_user_id, session_name=session_name).order_by(AmsMonitoring.auditee).all()
        
        data_list = []
        for auditee in auditees:
            data_list.append({
                "id": auditee.id,
                "auditee": auditee.auditee,
                "periode": auditee.periode.strftime('%Y-%m') if auditee.periode else "",
                "periode_display": auditee.periode.strftime('%B %Y') if auditee.periode else "-",
                "tahun_audit": auditee.tahun_audit,
                "total_rekomendasi": auditee.total_rekomendasi,
                "monitoring_type": auditee.monitoring_type, 
                "selesai": auditee.selesai,
                "tidak_selesai": auditee.tidak_selesai,
                "todo": auditee.todo,
                "belum_sesuai": auditee.belum_sesuai,
                "belum_tl": auditee.belum_tl,
                "tdd": auditee.tdd,
                "is_read_only": (str(target_user_id) != str(current_user.id))
            })
        return jsonify(data_list), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/add_auditee_data', methods=['POST'])
@login_required
def api_add_auditee_data():
    data = request.get_json()
    try:
        periode_str = data.get('periode')
        periode_date = datetime.date.today()
        if periode_str:
            periode_date = datetime.datetime.strptime(periode_str + '-01', '%Y-%m-%d').date()

        new_auditee = AmsMonitoring(
            user_id=current_user.id,
            session_name=data.get('session_name'),
            monitoring_type=data.get('monitoring_type'),
            periode=periode_date, 
            auditee=data.get('auditee'),
            tahun_audit=int(data.get('tahun_audit')),
            total_rekomendasi=int(data.get('total_rekomendasi', 0)),
            selesai=int(data.get('selesai', 0)),
            tidak_selesai=int(data.get('tidak_selesai', 0)),
            todo=int(data.get('todo', 0)),
            belum_sesuai=int(data.get('belum_sesuai', 0)),
            belum_tl=int(data.get('belum_tl', 0)),
            tdd=int(data.get('tdd', 0))
        )
        db.session.add(new_auditee)
        db.session.commit()
        return jsonify({"status": "success", "message": "Data berhasil ditambahkan."}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete_auditee/<int:auditee_id>', methods=['DELETE'])
@login_required
def api_delete_auditee(auditee_id):
    """Menghapus data auditee."""
    auditee = AmsMonitoring.query.get(auditee_id)
    if not auditee or auditee.user_id != current_user.id:
        return jsonify({"error": "Data tidak ditemukan atau akses ditolak."}), 404

    try:
        db.session.delete(auditee)
        db.session.commit()
        return jsonify({"status": "success", "message": "Data auditee berhasil dihapus."}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus auditee: {e}")
        return jsonify({"error": "Gagal menghapus data auditee."}), 500

@app.route('/ams_tl_tidak_setuju')
@login_required
def ams_tl_tidak_setuju_page():
    return render_template('ams_tl_tidak_setuju.html', username=current_user.username, label=current_user.label)

@app.route('/api/tl_tidak_setuju/add', methods=['POST'])
@login_required
def api_add_tl_tidak_setuju():
    data = request.json
    try:
        def parse_date(date_str):
            return datetime.datetime.strptime(date_str, '%Y-%m-%d').date() if date_str else None

        new_data = AmsTlTidakSetuju(
            user_id=current_user.id,
            no_aoi=data.get('no_aoi'),
            jenis_aoi=data.get('jenis_aoi'),
            klasifikasi=data.get('klasifikasi'),
            no_lha=data.get('no_lha'),
            nama_penugasan=data.get('nama_penugasan'),
            keterangan=data.get('keterangan'),
            temuan=data.get('temuan'),
            rekomendasi=data.get('rekomendasi'),
            auditee=data.get('auditee'),
            target_per_lha=parse_date(data.get('target_per_lha')),
            perubahan_target_date=parse_date(data.get('perubahan_target_date')),
            tindak_lanjut=data.get('tindak_lanjut')
        )
        db.session.add(new_data)
        db.session.commit()
        return jsonify({"status": "success", "message": "Data berhasil ditambahkan"}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/tl_tidak_setuju/get', methods=['GET'])
@login_required
def api_get_tl_tidak_setuju():
    my_data = AmsTlTidakSetuju.query.filter_by(user_id=current_user.id).all()
    shared_records = SharedTlTidakSetuju.query.filter_by(shared_with_id=current_user.id).all()
    owner_ids = [s.owner_id for s in shared_records]
    shared_data_list = []
    if owner_ids:
        shared_data_list = AmsTlTidakSetuju.query.filter(AmsTlTidakSetuju.user_id.in_(owner_ids)).all()
    all_items = my_data + shared_data_list
    
    result = []
    for item in all_items:
        is_shared_item = (item.user_id != current_user.id)
        owner_name = item.user.username if is_shared_item else "Me"

        result.append({
            "id": item.id,
            "no_aoi": item.no_aoi,
            "jenis_aoi": item.jenis_aoi,
            "klasifikasi": item.klasifikasi,
            "no_lha": item.no_lha,
            "nama_penugasan": item.nama_penugasan,
            "keterangan": item.keterangan,
            "temuan": item.temuan,
            "rekomendasi": item.rekomendasi,
            "auditee": item.auditee,
            "target_per_lha": item.target_per_lha.strftime('%Y-%m-%d') if item.target_per_lha else "",
            "perubahan_target_date": item.perubahan_target_date.strftime('%Y-%m-%d') if item.perubahan_target_date else "",
            "tindak_lanjut": item.tindak_lanjut,
            "is_shared": is_shared_item,
            "owner_name": owner_name
        })
    return jsonify(result), 200

@app.route('/api/tl_tidak_setuju/delete/<int:id>', methods=['DELETE'])
@login_required
def api_delete_tl_tidak_setuju(id):
    item = AmsTlTidakSetuju.query.get(id)
    if not item or item.user_id != current_user.id:
        return jsonify({"error": "Data tidak ditemukan"}), 404
    db.session.delete(item)
    db.session.commit()
    return jsonify({"status": "success"}), 200

@app.route('/api/tl_tidak_setuju/edit/<int:id>', methods=['POST'])
@login_required
def api_edit_tl_tidak_setuju(id):
    item = AmsTlTidakSetuju.query.get(id)
    if not item or item.user_id != current_user.id:
        return jsonify({"error": "Data tidak ditemukan"}), 404
    
    data = request.json
    def parse_date(date_str):
        return datetime.datetime.strptime(date_str, '%Y-%m-%d').date() if date_str else None

    try:
        item.no_aoi = data.get('no_aoi')
        item.jenis_aoi = data.get('jenis_aoi')
        item.klasifikasi = data.get('klasifikasi')
        item.no_lha = data.get('no_lha')
        item.nama_penugasan = data.get('nama_penugasan')
        item.keterangan = data.get('keterangan')
        item.temuan = data.get('temuan')
        item.rekomendasi = data.get('rekomendasi')
        item.auditee = data.get('auditee')
        item.target_per_lha = parse_date(data.get('target_per_lha'))
        item.perubahan_target_date = parse_date(data.get('perubahan_target_date'))
        item.tindak_lanjut = data.get('tindak_lanjut')
        
        db.session.commit()
        return jsonify({"status": "success", "message": "Data berhasil diupdate"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/tl_tidak_setuju/share', methods=['POST'])
@login_required
def api_share_tl_tidak_setuju():
    """Membagikan Akses TL Tidak Setuju ke user lain"""
    data = request.json
    target_user_ids = data.get('user_ids', [])

    if not target_user_ids:
        return jsonify({"error": "Tidak ada user yang dipilih."}), 400

    success_count = 0
    for uid in target_user_ids:
        try:
            exists = SharedTlTidakSetuju.query.filter_by(
                owner_id=current_user.id, 
                shared_with_id=uid
            ).first()

            if not exists:
                new_share = SharedTlTidakSetuju(
                    owner_id=current_user.id,
                    shared_with_id=uid
                )
                db.session.add(new_share)
                success_count += 1
        except:
            continue

    try:
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil dibagikan ke {success_count} pengguna."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/add_comment', methods=['POST'])
@login_required 
def add_comment():
    data = request.get_json()
    folder_name = data.get('folderName')
    filename = data.get('fileName')
    row_id = data.get('rowId')
    text = data.get('text')
    parent_id = data.get('parentId') 
    if not folder_name or not filename or not row_id or not text:
        return jsonify({'status': 'error', 'message': 'Data komentar tidak lengkap.'}), 400

    try:
        new_comment = Comment(
            owner_id=current_user.id,
            folder_name=folder_name,
            filename=filename,
            row_id=row_id,
            username=current_user.fullname,
            text=text,
            parent_id=parent_id
        )
        db.session.add(new_comment)
        db.session.commit()

        return jsonify({'status': 'success', 'message': 'Komentar berhasil disimpan.'}), 200

    except Exception as e:
        db.session.rollback()
        print('Error saat menyimpan komentar:', e)
        return jsonify({'status': 'error', 'message': f"Gagal menyimpan komentar: {str(e)}"}), 500
    
@app.route('/api/save_row_action', methods=['POST'])
@login_required
def api_save_row_action():
    """
    Menyimpan status checkbox dan PIC ke database.
    PERBAIKAN: Konversi tipe data yang ketat & Debugging print.
    """
    try:
        data = request.get_json()
        if not isinstance(data, dict):
            return jsonify({"error": "Format data salah."}), 400
            
        folder_name = data.get('folder_name')
        filename = data.get('filename')
        
        # PERBAIKAN: Konversi ID ke Integer
        try:
            owner_id = int(data.get('owner_id'))
            row_id = int(data.get('row_id'))
        except (TypeError, ValueError):
            return jsonify({"error": "ID tidak valid (harus angka)."}), 400

        is_ganti = data.get('is_ganti', False)
        
        # Handle PIC User ID (bisa string kosong atau angka)
        raw_pic = data.get('pic_user_id')
        pic_user_id = None
        if raw_pic is not None and str(raw_pic).strip() != "":
            if str(raw_pic).isdigit():
                pic_user_id = int(raw_pic)

        if not all([folder_name, filename, owner_id is not None, row_id is not None]):
            return jsonify({"error": "Data tidak lengkap."}), 400

        # DEBUGGING: Print ke terminal untuk cek data yang masuk
        print(f"[DEBUG SAVE] Saving Row: Owner={owner_id}, Folder={folder_name}, File={filename}, Row={row_id}, PIC={pic_user_id}")

        # Cek Akses
        is_owner = (current_user.id == owner_id)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak."}), 403

        action = RowAction.query.filter_by(
            owner_id=owner_id,
            folder_name=folder_name,
            filename=filename,
            row_id=row_id
        ).first()

        if action:
            action.is_ganti = is_ganti
            action.pic_user_id = pic_user_id
            print(f"[DEBUG SAVE] Updated existing row {row_id}")
        else:
            action = RowAction(
                owner_id=owner_id,
                folder_name=folder_name,
                filename=filename,
                row_id=row_id,
                is_ganti=is_ganti,
                pic_user_id=pic_user_id
            )
            db.session.add(action)
            print(f"[DEBUG SAVE] Created new row {row_id}")
        
        db.session.commit()
        return jsonify({"status": "success", "message": "Status tersimpan."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"[ERROR SAVE] {e}")
        return jsonify({"error": f"Server Error: {str(e)}"}), 500

@app.route('/ams_temuan')
@login_required
def ams_temuan_page():
    """Halaman utama fitur Temuan."""
    return render_template('ams_temuan.html', 
                           username=current_user.username, 
                           label=current_user.label)

@app.route('/api/get_temuan_sessions', methods=['GET'])
@login_required
def api_get_temuan_sessions():
    """Mengambil daftar sesi temuan milik user DAN sesi yang dishare ke user."""
    try:
        own_sessions = AmsTemuanSession.query.filter_by(user_id=current_user.id).order_by(AmsTemuanSession.created_at.desc()).all()
        shared_records = SharedTemuanSession.query.filter_by(shared_with_id=current_user.id).all()
        shared_session_ids = [s.session_id for s in shared_records]
        shared_sessions = []
        if shared_session_ids:
            shared_sessions = AmsTemuanSession.query.filter(AmsTemuanSession.id.in_(shared_session_ids)).all()
        data = []

        for s in own_sessions:
            data.append({
                "id": s.id,
                "nama_sesi": s.nama_sesi,
                "jenis_audit": s.jenis_audit,
                "created_at": s.created_at.strftime('%Y-%m-%d'),
                "is_owner": True
            })
            
        for s in shared_sessions:
            owner_name = s.user.fullname if s.user else "Unknown"
            data.append({
                "id": s.id,
                "nama_sesi": f"{s.nama_sesi} (Oleh: {owner_name})", 
                "jenis_audit": s.jenis_audit,
                "created_at": s.created_at.strftime('%Y-%m-%d'),
                "is_owner": False
            })
            
        return jsonify(data), 200
    except Exception as e:
        print(f"Error get temuan sessions: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/add_temuan_session', methods=['POST'])
@login_required
def api_add_temuan_session():
    """Membuat sesi temuan baru."""
    data = request.json
    try:
        new_session = AmsTemuanSession(
            user_id=current_user.id,
            nama_sesi=data.get('session_name'),
            jenis_audit=data.get('jenis_audit')
        )
        db.session.add(new_session)
        db.session.commit()
        return jsonify({"status": "success", "id": new_session.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete_temuan_session/<int:id>', methods=['DELETE'])
@login_required
def api_delete_temuan_session(id):
    """Menghapus sesi temuan (Cascade delete ke rows otomatis)."""
    session = AmsTemuanSession.query.get(id)
    if not session or session.user_id != current_user.id:
        return jsonify({"error": "Sesi tidak ditemukan"}), 404
    try:
        db.session.delete(session)
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/share_temuan_session', methods=['POST'])
@login_required
def api_share_temuan_session():
    """Membagikan sesi Temuan ke user lain."""
    data = request.json
    session_id = data.get('session_id') 
    target_user_ids = data.get('user_ids', [])

    if not session_id or not target_user_ids:
        return jsonify({"error": "Data tidak lengkap."}), 400

    session = AmsTemuanSession.query.get(session_id)
    if not session:
        return jsonify({"error": "Sesi tidak ditemukan."}), 404
        
    if session.user_id != current_user.id:
        return jsonify({"error": "AKSES DITOLAK: Hanya pembuat sesi yang boleh membagikan sesi ini."}), 403

    success_count = 0
    for uid in target_user_ids:
        try:
            if int(uid) == current_user.id: continue
            exists = SharedTemuanSession.query.filter_by(
                owner_id=current_user.id,
                shared_with_id=uid,
                session_id=session_id
            ).first()

            if not exists:
                new_share = SharedTemuanSession(
                    owner_id=current_user.id,
                    shared_with_id=uid,
                    session_id=session_id
                )
                db.session.add(new_share)
                success_count += 1
        except Exception:
            continue

    try:
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil dibagikan ke {success_count} pengguna."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_temuan_data/<int:session_id>', methods=['GET'])
@login_required
def api_get_temuan_data(session_id):
    """
    Mengambil semua baris data dalam satu sesi, 
    LENGKAP dengan list komentar dan history edit untuk setiap baris.
    """
    session = AmsTemuanSession.query.get(session_id)
    if not session:
        return jsonify({"error": "Sesi tidak ditemukan"}), 404
    is_owner = (session.user_id == current_user.id)
    is_shared = False
    if not is_owner:
        share_check = SharedTemuanSession.query.filter_by(
            session_id=session_id, 
            shared_with_id=current_user.id
        ).first()
        if share_check:
            is_shared = True

    if not is_owner and not is_shared:
        return jsonify({"error": "Akses ditolak ke sesi ini"}), 403
    
    rows = AmsTemuanRow.query.filter_by(session_id=session_id).all()
    
    data_list = []
    
    for r in rows:
        comments = AmsTemuanComment.query.filter_by(row_id=r.id).order_by(AmsTemuanComment.created_at.asc()).all()
        
        comments_data = []
        for c in comments:
            comments_data.append({
                "id": c.id,
                "username": c.username,
                "content": c.content,
                "created_at": c.created_at.strftime("%d %b %Y, %H:%M"),
                "is_owner": (c.user_id == current_user.id) 
            })

        history_records = AmsRowHistory.query.filter_by(row_id=r.id).order_by(AmsRowHistory.timestamp.asc()).all()
        history_list = []
        for h in history_records:
            history_list.append({
                "user": h.username,
                "time": h.timestamp.strftime("%d-%m-%Y %H:%M"), 
                "columns": h.changed_columns 
            })

        data_list.append({
            "id": r.id,
            "pic_skai": r.pic_skai,
            "no_aoi": r.no_aoi,
            "jenis_aoi": r.jenis_aoi,
            "klasifikasi": r.klasifikasi,
            "no_lha": r.no_lha,
            "nama_penugasan": r.nama_penugasan,
            "aoi": r.aoi,
            "rekomendasi": r.rekomendasi,
            "rencana_tl": r.rencana_tl,
            "rencana_evidence": r.rencana_evidence,
            "auditee": r.auditee,
            "pic_auditee": r.pic_auditee,
            "target_penyelesaian": r.target_penyelesaian, 
            "perubahan_target": r.perubahan_target,
            "tindak_lanjut": r.tindak_lanjut,
            "signifikansi": r.signifikansi,
            "jml_rekomendasi": r.jml_rekomendasi,
            "selesai": r.selesai,
            "belum_jt_bs": r.belum_jt_bs,
            "os_bd": r.os_bd,
            "tdd": r.tdd,
            "control": r.control,
            "history_logs": history_list,
            "comments": comments_data,
            "last_modified_by": r.last_modified_by,
            "last_modified_at": r.last_modified_at.strftime("%d-%m-%Y %H:%M") if r.last_modified_at else "-"
        })

    return jsonify(data_list), 200

@app.route('/api/add_temuan_row', methods=['POST'])
@login_required
def api_add_temuan_row():
    """Menambahkan satu baris data ke sesi temuan (Support Share)."""
    data = request.json
    session_id = data.get('session_id')
    
    session = AmsTemuanSession.query.get(session_id)
    if not session:
        return jsonify({"error": "Sesi tidak ditemukan"}), 404
    is_owner = (session.user_id == current_user.id)
    is_shared = False
    if not is_owner:
        if SharedTemuanSession.query.filter_by(session_id=session_id, shared_with_id=current_user.id).first():
            is_shared = True
    
    if not is_owner and not is_shared:
        return jsonify({"error": "Akses ditolak"}), 403
    
    def parse_int(v):
        try: return int(v)
        except: return 0

    try:
        new_row = AmsTemuanRow(
            session_id=session_id,
            user_id=current_user.id,
            pic_skai=data.get('pic_skai'),
            no_aoi=data.get('no_aoi'),
            jenis_aoi=data.get('jenis_aoi'),
            klasifikasi=data.get('klasifikasi'),
            no_lha=data.get('no_lha'),
            nama_penugasan=data.get('nama_penugasan'),
            aoi=data.get('aoi'),
            rekomendasi=data.get('rekomendasi'),
            rencana_tl=data.get('rencana_tl'),
            rencana_evidence=data.get('rencana_evidence'),
            auditee=data.get('auditee'),
            pic_auditee=data.get('pic_auditee'),
            target_penyelesaian=data.get('target_penyelesaian'),
            perubahan_target=data.get('perubahan_target'),
            tindak_lanjut=data.get('tindak_lanjut'),
            signifikansi=data.get('signifikansi'),
            jml_rekomendasi=parse_int(data.get('jml_rekomendasi')),
            selesai=parse_int(data.get('selesai')),
            belum_jt_bs=parse_int(data.get('belum_jt_bs')),
            os_bd=parse_int(data.get('os_bd')),
            tdd=parse_int(data.get('tdd')),
            control=data.get('control'),
            last_modified_by=current_user.fullname,
            last_modified_at=datetime.datetime.now()
        )
        
        db.session.add(new_row)
        db.session.commit()
        return jsonify({"status": "success", "message": "Data berhasil disimpan"}), 201
    except Exception as e:
        db.session.rollback()
        print(f"Error adding row: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete_temuan_row/<int:id>', methods=['DELETE'])
@login_required
def api_delete_temuan_row(id):
    """Menghapus satu baris data (Support Share)."""
    row = AmsTemuanRow.query.get(id)
    if not row:
        return jsonify({"error": "Data tidak ditemukan"}), 404
    
    is_owner = (row.user_id == current_user.id)
    session = AmsTemuanSession.query.get(row.session_id)
    is_session_owner = (session.user_id == current_user.id)
    is_shared = False

    if not is_owner and not is_session_owner:
        if SharedTemuanSession.query.filter_by(session_id=row.session_id, shared_with_id=current_user.id).first():
            is_shared = True

    if not is_owner and not is_session_owner and not is_shared:
        return jsonify({"error": "Akses ditolak"}), 403
    try:
        db.session.delete(row)
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/log_analysis')
@login_required
def log_analysis_page():
    """Menampilkan halaman log analisis dan tracker tugas."""
    return render_template('log_analysis.html', username=current_user.username, label=current_user.label)

@app.route('/api/log_analysis_start', methods=['POST'])
@login_required
def api_log_analysis_start():
    """Mencatat awal proses analisis."""
    data = request.json
    filename = data.get('filename')
    feature_type = data.get('feature_type')

    if not filename or not feature_type:
        return jsonify({"error": "Data tidak lengkap."}), 400

    try:
        new_log = AnalysisLog(
            user_id=current_user.id,
            filename=filename,
            feature_type=feature_type,
            status='unfinished'
        )
        db.session.add(new_log)
        db.session.commit()

        return jsonify({
            "status": "success",
            "log_id": new_log.id
        }), 201
    except Exception as e:
        db.session.rollback()
        print(f"Error saat memulai log analisis: {e}")
        return jsonify({"error": "Gagal memulai log."}), 500

@app.route('/api/log_analysis_end', methods=['POST'])
@login_required
def api_log_analysis_end():
    """Mencatat akhir proses analisis."""
    data = request.json
    log_id = data.get('log_id')
    status = data.get('status') 

    if not log_id or status not in ['done', 'error']:
        return jsonify({"error": "Data tidak lengkap atau status tidak valid."}), 400

    try:
        log_entry = AnalysisLog.query.get(log_id)
        if not log_entry or log_entry.user_id != current_user.id:
            return jsonify({"error": "Log tidak ditemukan atau akses ditolak."}), 404
        
        log_entry.end_time = datetime.datetime.utcnow()
        log_entry.status = status
        db.session.commit()
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat mengakhiri log analisis: {e}")
        return jsonify({"error": "Gagal mengakhiri log."}), 500

@app.route('/api/tl_tidak_setuju/upload_image', methods=['POST'])
@login_required
def api_upload_tl_image():
    """OCR Gambar Tabel via Gemini (LOGIKA: APPEND / MENAMBAHKAN)"""
    if 'file' not in request.files: return jsonify({"error": "File tidak ditemukan"}), 400
    file = request.files['file']
    
    try:
        img = PIL.Image.open(file)
        prompt = """
        Analisis gambar tabel ini. Ekstrak datanya baris per baris.
        Abaikan header. Petakan ke format JSON object dengan key berikut:
        - no_aoi (String)
        - jenis_aoi (String)
        - klasifikasi (String)
        - no_lha (String)
        - keterangan (String)
        - temuan (String)
        - rekomendasi (String)
        - auditee (String)
        - target_per_lha (Format: YYYY-MM-DD, jika tidak ada null)
        - perubahan_target_date (Format: YYYY-MM-DD, jika tidak ada null)
        - tindak_lanjut (String)

        PENTING: Output HANYA JSON Array. Jangan pakai markdown.
        """
        
        model = genai.GenerativeModel('gemini-2.5-pro')
        track_gemini_usage()
        response = model.generate_content([prompt, img])
        cleaned = response.text.replace('```json', '').replace('```', '').strip()
        data_list = json.loads(cleaned)

        count = 0
        for item in data_list:
            def parse_date(d): 
                try:
                    return datetime.datetime.strptime(d, '%Y-%m-%d').date() if d else None
                except:
                    return None

            new_data = AmsTlTidakSetuju(
                user_id=current_user.id,
                no_aoi=item.get('no_aoi'),
                jenis_aoi=item.get('jenis_aoi'),
                klasifikasi=item.get('klasifikasi'),
                no_lha=item.get('no_lha'),
                nama_penugasan=item.get('nama_penugasan'),
                keterangan=item.get('keterangan'),
                temuan=item.get('temuan'),
                rekomendasi=item.get('rekomendasi'),
                auditee=item.get('auditee'),
                target_per_lha=parse_date(item.get('target_per_lha')),
                perubahan_target_date=parse_date(item.get('perubahan_target_date')),
                tindak_lanjut=item.get('tindak_lanjut')
            )
            db.session.add(new_data) 
            count += 1
        
        db.session.commit()

        total_now = AmsTlTidakSetuju.query.filter_by(user_id=current_user.id).count()

        return jsonify({
            "status": "success", 
            "message": f"Berhasil MENAMBAHKAN {count} data baru. Total data Anda sekarang: {total_now}"
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Gagal proses gambar: {str(e)}"}), 500

@app.route('/api/tl_tidak_setuju/upload_excel', methods=['POST'])
@login_required
def api_upload_tl_excel():
    """Import dari Excel"""
    if 'file' not in request.files: return jsonify({"error": "File tidak ditemukan"}), 400
    file = request.files['file']

    try:
        df = pd.read_excel(file)
        count = 0
        for _, row in df.iterrows():
            def get_val(col_name):
                val = row.get(col_name)
                return str(val) if pd.notna(val) else ""
            
            def get_date(col_name):
                val = row.get(col_name)
                if pd.notna(val):
                    try:
                        return pd.to_datetime(val).date()
                    except:
                        return None
                return None

            new_data = AmsTlTidakSetuju(
                user_id=current_user.id,
                no_aoi=get_val('No AOI') or get_val('no_aoi'),
                jenis_aoi=get_val('Jenis AOI') or get_val('jenis_aoi'),
                klasifikasi=get_val('Klasifikasi') or get_val('klasifikasi'),
                no_lha=get_val('No LHA') or get_val('no_lha'),
                keterangan=get_val('Keterangan') or get_val('keterangan'),
                temuan=get_val('Temuan') or get_val('temuan'),
                rekomendasi=get_val('Rekomendasi') or get_val('rekomendasi'),
                auditee=get_val('Auditee') or get_val('auditee'),
                target_per_lha=get_date('Target per LHA') or get_date('target_per_lha'),
                perubahan_target_date=get_date('Perubahan Target Date') or get_date('perubahan_target_date'),
                tindak_lanjut=get_val('Tindak Lanjut') or get_val('tindak_lanjut')
            )
            db.session.add(new_data)
            count += 1
            
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil import {count} baris dari Excel."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Gagal proses Excel: {str(e)}"}), 500

@app.route('/api/import_temuan_json', methods=['POST'])
@login_required
def api_import_temuan_json():
    try:
        req = request.get_json()
        session_id = req.get('session_id')
        data_list = req.get('data')
        should_reset = req.get('reset_data', False)

        if not session_id or not data_list:
            return jsonify({"error": "Data tidak lengkap"}), 400

        if should_reset:
            AmsTemuanRow.query.filter_by(session_id=session_id).delete()
            db.session.commit()

        count = 0
        
        def parse_mixed_date(val):
            if not val: return None
            val_str = str(val).strip()
            if not val_str or val_str.lower() in ['-', 'nan', 'null', 'none', '0', '']: return None
            import re
            if re.match(r'^\d+(\.\d+)?$', val_str):
                try:
                    days = float(val_str)
                    return (datetime.datetime(1899, 12, 30) + datetime.timedelta(days=days)).strftime('%Y-%m-%d')
                except: pass 
            formats = ['%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y', '%d/%m/%Y', '%Y/%m/%d', '%d-%b-%y', '%d %b %Y', '%B %d, %Y', '%d-%b-%Y', '%d %B %Y', '%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M:%S.%f']
            for fmt in formats:
                try: return datetime.datetime.strptime(val_str, fmt).strftime('%Y-%m-%d')
                except ValueError: continue
            return val_str

        def parse_safe_int(val):
            if val is None or val == "": return 0
            val_str = str(val).strip().replace(',', '.') 
            try:
                import re
                numeric_part = re.search(r'-?\d+\.?\d*', val_str)
                return int(float(numeric_part.group())) if numeric_part else 0
            except: return 0

        def get_val(row, keywords):
            row_normalized = {}
            for k, v in row.items():
                clean_key = str(k).lower().replace('.', '').replace(' ', '').replace('_', '')
                row_normalized[clean_key] = v
            
            for k in keywords:
                clean_keyword = k.lower().replace(' ', '').replace('_', '').replace('.', '')
                if clean_keyword in row_normalized:
                    val = row_normalized[clean_keyword]
                    return str(val) if val is not None else ""
            return ""

        last_no_aoi = ""
        last_auditee = ""

        for row in data_list:
            raw_no_aoi = get_val(row, ['noaoi', 'nomor', 'nomoraoi'])
            raw_auditee = get_val(row, ['auditee', 'unit'])

            current_no_aoi = raw_no_aoi if raw_no_aoi else last_no_aoi
            current_auditee = raw_auditee if raw_auditee else last_auditee

            if raw_no_aoi: last_no_aoi = raw_no_aoi
            if raw_auditee: last_auditee = raw_auditee

            target_val = parse_mixed_date(get_val(row, ['target', 'targetpenyelesaian', 'duedate', 'tanggaltarget']))
            perubahan_val = parse_mixed_date(get_val(row, ['perubahan', 'perubahantarget', 'perubahantargetpenyelesaian', 'reviseddate']))

            new_row = AmsTemuanRow(
                session_id=session_id,
                user_id=current_user.id,
                pic_skai=str(get_val(row, ['picskai', 'pic', 'pic_skai'])), 
                no_aoi=current_no_aoi, 
                jenis_aoi=str(get_val(row, ['jenis', 'jenisaoi'])),
                klasifikasi=str(get_val(row, ['klasifikasi', 'tahun'])),
                no_lha=str(get_val(row, ['nolha', 'nomorlha', 'lha'])),
                nama_penugasan=str(get_val(row, ['nama_penugasan', 'namapenugasan', 'nama', 'penugasan', 'judul'])),
                aoi=str(get_val(row, ['aoi', 'areaofimprovement', 'temuan'])),
                rekomendasi=str(get_val(row, ['rekomendasi'])),
                rencana_tl=str(get_val(row, ['rencanatl', 'rencanatindaklanjut'])),
                rencana_evidence=str(get_val(row, ['evidence', 'rencanaevidence'])),
                auditee=current_auditee, 
                pic_auditee=str(get_val(row, ['pic_auditee', 'picauditee'])), 
                target_penyelesaian=target_val,
                perubahan_target=perubahan_val,
                tindak_lanjut=str(get_val(row, ['tindaklanjut', 'tindak_lanjut', 'tl', 'progress', 'status'])),
                signifikansi=str(get_val(row, ['signifikansi'])),
                jml_rekomendasi=parse_safe_int(get_val(row, ['jml_rekomendasi', 'jmlrekomendasi', 'jumlah', 'jml', 'total'])),
                selesai=parse_safe_int(get_val(row, ['selesai', 'done'])),
                belum_jt_bs=parse_safe_int(get_val(row, ['belum_jt_bs', 'belumjtbs', 'bjt', 'bs', 'belum_jatuh_tempo_bs_bpk'])),
                os_bd=parse_safe_int(get_val(row, ['os_bd', 'osbd', 'os', 'bd', 'os_bd_bpk'])),
                tdd=parse_safe_int(get_val(row, ['tdd', 'tdd_bpk'])),
                
                control=str(get_val(row, ['control'])),
                last_modified_by=current_user.fullname,
                last_modified_at=datetime.datetime.now()
            )
            db.session.add(new_row)
            count += 1
        
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil mengimpor {count} data."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Import Error: {e}")
        return jsonify({"error": f"Server Error: {str(e)}"}), 500

@app.route('/api/add_reminder', methods=['POST'])
@login_required
def api_add_reminder():
    data = request.json
    try:
        deadline_date = datetime.datetime.strptime(data['deadline'], '%Y-%m-%d').date()
        def list_to_str(lst):
            return ", ".join(lst) if isinstance(lst, list) else str(lst)

        new_reminder = AmsReminder(
            user_id=current_user.id,
            assigned_to=data['auditee'],
            pic_skai=list_to_str(data.get('pic_skai', [])),
            pic_reminder=list_to_str(data.get('pic_reminder', [])),
            pic_auditee=list_to_str(data.get('pic_auditee', [])),
            tembusan=data.get('tembusan', ''),
            subject=data['subject'],
            temuan=data.get('temuan', ''),
            deadline=deadline_date
        )
        db.session.add(new_reminder)
        db.session.commit()
        return jsonify({"status": "success", "message": "Reminder berhasil ditambahkan."}), 201
    except Exception as e:
        db.session.rollback()
        print(f"Error adding reminder: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/check_deadline_notifications', methods=['GET'])
@login_required
def api_check_deadline_notifications():
    try:
        today = datetime.date.today()
        tomorrow = today + datetime.timedelta(days=1)
        reminders = AmsReminder.query.filter(
            AmsReminder.deadline == tomorrow,
            AmsReminder.pic_skai.contains(current_user.username),
            AmsReminder.is_reminded == False,  
            AmsReminder.is_responded == False  
        ).all()

        notifications = []
        for rem in reminders:
            notifications.append({
                "id": rem.id,
                "auditee": rem.assigned_to,
                "subject": rem.subject,
                "deadline": rem.deadline.strftime('%d-%m-%Y')
            })
            
        return jsonify(notifications), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/edit_temuan_row/<int:id>', methods=['POST'])
@login_required
def api_edit_temuan_row(id):
    row = AmsTemuanRow.query.get(id)
    if not row:
        return jsonify({"error": "Data tidak ditemukan"}), 404

    is_owner = (row.user_id == current_user.id)
    is_shared = False
    if not is_owner:
        if SharedTemuanSession.query.filter_by(session_id=row.session_id, shared_with_id=current_user.id).first():
            is_shared = True
    
    if not is_owner and not is_shared:
        return jsonify({"error": "Akses ditolak."}), 403

    data = request.json
    
    def parse_date(d): 
        try: return datetime.datetime.strptime(d, '%Y-%m-%d').date() if d else None
        except: return None
    def parse_int(v):
        try: return int(v)
        except: return 0

    field_mapping = {
        'no_aoi': 'No AOI', 'jenis_aoi': 'Jenis AOI', 'klasifikasi': 'Klasifikasi',
        'no_lha': 'No LHA', 'nama_penugasan': 'Nama Penugasan', 'aoi': 'AOI',
        'rekomendasi': 'Rekomendasi', 'rencana_tl': 'Rencana TL',
        'rencana_evidence': 'Rencana Evidence', 'auditee': 'Auditee',
        'pic_auditee': 'PIC Auditee', 'target_penyelesaian': 'Target Penyelesaian',
        'perubahan_target': 'Perubahan Target', 'tindak_lanjut': 'Tindak Lanjut',
        'signifikansi': 'Signifikansi', 'jml_rekomendasi': 'Jml Rekomendasi',
        'selesai': 'Selesai', 'belum_jt_bs': 'Belum JT/BS', 'os_bd': 'OS/BD',
        'tdd': 'TDD', 'control': 'Control'
    }

    changes_found = []

    for field, label in field_mapping.items():
        if field in ['jml_rekomendasi', 'selesai', 'belum_jt_bs', 'os_bd', 'tdd', 'target_penyelesaian', 'perubahan_target']:
            continue
            
        old_val = str(getattr(row, field) or '').strip()
        new_val = str(data.get(field) or '').strip()
        
        if old_val != new_val:
            changes_found.append(label)

    int_fields = ['jml_rekomendasi', 'selesai', 'belum_jt_bs', 'os_bd', 'tdd']
    for f in int_fields:
        old_val = int(getattr(row, f) or 0)
        new_val = parse_int(data.get(f))
        if old_val != new_val:
            changes_found.append(field_mapping[f])

    date_fields = ['target_penyelesaian', 'perubahan_target']
    for f in date_fields:
        old_val = str(getattr(row, f) or '')
        new_val = str(data.get(f) or '')
        if old_val != new_val:
            changes_found.append(field_mapping[f])

    try:
        row.pic_skai = data.get('pic_skai')
        row.no_aoi = data.get('no_aoi')
        row.jenis_aoi = data.get('jenis_aoi')
        row.klasifikasi = data.get('klasifikasi')
        row.no_lha = data.get('no_lha')
        row.nama_penugasan = data.get('nama_penugasan')
        row.aoi = data.get('aoi')
        row.rekomendasi = data.get('rekomendasi')
        row.rencana_tl = data.get('rencana_tl')
        row.rencana_evidence = data.get('rencana_evidence')
        row.auditee = data.get('auditee')
        row.pic_auditee = data.get('pic_auditee')
        row.target_penyelesaian = data.get('target_penyelesaian')
        row.perubahan_target = data.get('perubahan_target')
        row.tindak_lanjut = data.get('tindak_lanjut')
        row.signifikansi = data.get('signifikansi')
        row.jml_rekomendasi = parse_int(data.get('jml_rekomendasi'))
        row.selesai = parse_int(data.get('selesai'))
        row.belum_jt_bs = parse_int(data.get('belum_jt_bs'))
        row.os_bd = parse_int(data.get('os_bd'))
        row.tdd = parse_int(data.get('tdd'))
        row.control = data.get('control')

        if changes_found:
            cols_str = ", ".join(changes_found)
            
            history_entry = AmsRowHistory(
                row_id=row.id,
                user_id=current_user.id,
                username=current_user.fullname,
                changed_columns=cols_str,
                timestamp=datetime.datetime.now()
            )
            db.session.add(history_entry)

        db.session.commit()
        return jsonify({"status": "success", "message": "Data berhasil diperbarui"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/add_comment/<int:row_id>', methods=['POST'])
@login_required
def api_add_comment(row_id):
    data = request.json

    raw_content = data.get('comment', '')
    content = raw_content.strip() 
    
    if not content:
        return jsonify({"error": "Komentar tidak boleh kosong"}), 400

    try:
        new_comment = AmsTemuanComment(
            row_id=row_id,
            user_id=current_user.id,
            username=current_user.fullname,
            content=content 
        )
        
        db.session.add(new_comment)
        db.session.commit()
        return jsonify({"status": "success", "message": "Komentar berhasil ditambahkan"}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error add comment: {e}")
        return jsonify({"error": str(e)}), 500
    
@app.route('/api/delete_comment/<int:comment_id>', methods=['DELETE'])
@login_required
def api_delete_comment(comment_id):
    """API BARU: Menghapus komentar spesifik berdasarkan ID."""
    comment = AmsTemuanComment.query.get(comment_id)
    
    if not comment:
        return jsonify({"error": "Komentar tidak ditemukan"}), 404

    if comment.user_id != current_user.id:
        return jsonify({"error": "Anda tidak memiliki izin menghapus komentar ini."}), 403

    try:
        db.session.delete(comment)
        db.session.commit()
        return jsonify({"status": "success", "message": "Komentar berhasil dihapus"}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error delete comment: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_reminders', methods=['GET'])
@login_required
def api_get_reminders():
    try:
        reminders = AmsReminder.query.filter(
            or_(
                AmsReminder.user_id == current_user.id,
                AmsReminder.pic_skai.contains(current_user.username),
                AmsReminder.pic_reminder.contains(current_user.username),
                AmsReminder.pic_auditee.contains(current_user.username)
            )
        ).order_by(AmsReminder.deadline.asc()).all()
        
        result = []
        today = datetime.date.today()
        
        for rem in reminders:
            creator_name = rem.user.fullname if rem.user else "Unknown"
            delta = (rem.deadline - today).days
            if delta < 0:
                status_text = "Sudah lewat deadline"
                sisa_hari = f"+{abs(delta)} hari"
                is_overdue = True
            elif delta == 0:
                status_text = "Deadline hari ini"
                sisa_hari = "0 hari"
                is_overdue = False
            else:
                status_text = "Belum lewat deadline"
                sisa_hari = f"{delta} hari"
                is_overdue = False

            in_skai = (current_user.username in rem.pic_skai) if rem.pic_skai else False
            in_reminder = (current_user.username in rem.pic_reminder) if rem.pic_reminder else False
            in_auditee = (current_user.username in rem.pic_auditee) if rem.pic_auditee else False
            is_pic = in_skai or in_reminder or in_auditee
            is_owner = (rem.user_id == current_user.id)
            visible_count = 0
            total_count = 0
            
            if rem.linked_temuan_ids:
                ids_list = [int(x) for x in rem.linked_temuan_ids.split(',') if x.isdigit()]
                linked_rows = AmsTemuanRow.query.filter(AmsTemuanRow.id.in_(ids_list)).all()
                total_count = len(linked_rows)
                is_viewer_admin = (is_owner or rem.linked_by == current_user.username or in_reminder)

                if is_viewer_admin:
                    visible_count = total_count
                else:
                    personal_rows = 0
                    for r in linked_rows:
                        if r.pic_skai and current_user.username in r.pic_skai:
                            personal_rows += 1
                    visible_count = personal_rows

            result.append({
                "id": rem.id,
                "auditee": rem.assigned_to,
                "pic_skai": rem.pic_skai,
                "pic_reminder": rem.pic_reminder,
                "pic_auditee": rem.pic_auditee,
                "temuan": rem.temuan,
                "creator_name": creator_name,
                "created_at": rem.created_at.strftime('%d %b %Y, %H:%M'),
                "subject": rem.subject,
                "deadline": rem.deadline.strftime('%d-%m-%Y'),
                "deadline_raw": rem.deadline.strftime('%Y-%m-%d'),
                "status": status_text,
                "sisa_hari": sisa_hari,
                "is_overdue": is_overdue,
                "is_responded": rem.is_responded,
                "is_reminded": rem.is_reminded,
                "linked_temuan_ids": rem.linked_temuan_ids,
                "visible_temuan_count": visible_count, 
                
                "is_current_user_pic": is_pic,
                "is_owner": is_owner
            })
            
        return jsonify(result), 200

    except Exception as e:
        print(f"Error getting reminders: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_dashboard_logs', methods=['GET'])
@login_required
def api_get_dashboard_logs():
    """Mengambil ringkasan Log Analisis untuk Widget Dashboard"""
    try:
        all_logs = AnalysisLog.query.filter_by(user_id=current_user.id).all()
        
        on_progress_list = []
        overdue_list = []
        
        today = datetime.date.today()

        for log in all_logs:
            if log.status == 'done':
                continue
            current_status = log.status
            if log.deadline and log.deadline.date() < today and log.end_time is None:
                current_status = 'overdue'
            log_data = {
                "id": log.id,
                "filename": log.filename,
                "deadline": log.deadline.strftime('%d %b %Y') if log.deadline else "Tanpa Deadline",
                "days_left": (log.deadline.date() - today).days if log.deadline else 999
            }

            if current_status == 'overdue':
                overdue_list.append(log_data)
            elif current_status == 'on_progress' or current_status == 'unfinished':
                on_progress_list.append(log_data)

        return jsonify({
            "on_progress": on_progress_list,
            "overdue": overdue_list
        }), 200

    except Exception as e:
        print(f"Error dashboard logs: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/edit_reminder/<int:id>', methods=['POST'])
@login_required
def api_edit_reminder(id):
    data = request.json
    reminder = AmsReminder.query.get(id)
    if not reminder:
        return jsonify({"error": "Reminder tidak ditemukan."}), 404

    is_owner = (reminder.user_id == current_user.id)
    is_pic_skai = (current_user.username in reminder.pic_skai) if reminder.pic_skai else False
    is_pic_reminder = (current_user.username in reminder.pic_reminder) if reminder.pic_reminder else False
    
    if not is_owner and not is_pic_skai and not is_pic_reminder:
        return jsonify({"error": "Akses ditolak. Hanya Owner atau PIC yang boleh mengedit."}), 403
    
    try:
        def list_to_str(lst):
            return ", ".join(lst) if isinstance(lst, list) else str(lst)

        reminder.assigned_to = data['auditee']
        reminder.subject = data['subject']
        reminder.deadline = datetime.datetime.strptime(data['deadline'], '%Y-%m-%d').date()
        reminder.pic_skai = list_to_str(data.get('pic_skai', []))
        reminder.pic_reminder = list_to_str(data.get('pic_reminder', []))
        reminder.pic_auditee = list_to_str(data.get('pic_auditee', [])) 
        reminder.tembusan = data.get('tembusan', '')
        
        if 'temuan' in data:
            reminder.temuan = data.get('temuan', '')

        db.session.commit()
        return jsonify({"status": "success", "message": "Reminder berhasil diperbarui."}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error editing reminder: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete_reminder/<int:id>', methods=['DELETE'])
@login_required
def api_delete_reminder(id):
    reminder = AmsReminder.query.get(id)

    if not reminder:
        return jsonify({"error": "Reminder tidak ditemukan."}), 404
    
    is_owner = (reminder.user_id == current_user.id)
    is_pic_reminder = False
    
    if reminder.pic_reminder:
        list_pic = [p.strip() for p in reminder.pic_reminder.split(',') if p.strip()]
        
        print(f"[DEBUG DELETE] Reminder ID: {id}")
        print(f"[DEBUG DELETE] List PIC di DB: {list_pic}")
        print(f"[DEBUG DELETE] Current User: {current_user.username} | Fullname: {current_user.fullname}")

        if current_user.username in list_pic or current_user.fullname in list_pic:
            is_pic_reminder = True

    if not is_owner and not is_pic_reminder:
        return jsonify({"error": "Akses ditolak. Hanya Owner atau PIC Reminder yang dapat menghapus."}), 403

    try:
        reminder.linked_temuan_ids = None 
        
        db.session.delete(reminder)
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus reminder: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/tl_tidak_setuju/import_json', methods=['POST'])
@login_required
def api_import_tl_json():
    try:
        req = request.get_json()
        data_list = req.get('data')

        if not data_list:
            return jsonify({"error": "Data kosong"}), 400

        count = 0
        
        def parse_date(val):
            if not val: return None
            val_str = str(val).strip()
            for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d', '%d-%b-%y'):
                try: return datetime.datetime.strptime(val_str, fmt).date()
                except ValueError: continue
            return None

        def normalize_key(row, keys_to_find):
            normalized_row = {}
            for k, v in row.items():
                clean_k = str(k).lower().replace(' ', '_').replace('.', '')
                normalized_row[clean_k] = v

            for key in keys_to_find:
                if key in normalized_row:
                    return normalized_row[key]
            return ''

        for row in data_list:
            new_data = AmsTlTidakSetuju(
                user_id=current_user.id,
                no_aoi=str(normalize_key(row, ['no_aoi', 'noaoi', 'nomor_aoi'])),
                jenis_aoi=str(normalize_key(row, ['jenis_aoi', 'jenis'])),
                klasifikasi=str(normalize_key(row, ['klasifikasi', 'class'])),
                no_lha=str(normalize_key(row, ['no_lha', 'nolha', 'lha'])),
                nama_penugasan=str(normalize_key(row, ['nama_penugasan', 'namapenugasan', 'judul'])),
                keterangan=str(normalize_key(row, ['keterangan', 'ket'])),
                temuan=str(normalize_key(row, ['temuan', 'isi_temuan'])),
                rekomendasi=str(normalize_key(row, ['rekomendasi', 'rek'])),
                auditee=str(normalize_key(row, ['auditee', 'unit'])),
                target_per_lha=parse_date(normalize_key(row, ['target_per_lha', 'target_lha'])),
                perubahan_target_date=parse_date(normalize_key(row, ['perubahan_target_date', 'perubahan_target'])),
                
                tindak_lanjut=str(normalize_key(row, ['tindak_lanjut', 'tl']))
            )
            db.session.add(new_data)
            count += 1

        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil mengimpor {count} data."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error Import: {e}")
        return jsonify({"error": f"Gagal Import: {str(e)}"}), 500

@app.route('/api/add_manual_task', methods=['POST'])
@login_required
def api_add_manual_task():
    """Menambahkan tugas manual ke log."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "Request body tidak valid (harus berupa JSON)."}), 400

    filename = data.get('filename')
    feature_type = data.get('feature_type')
    start_date_str = data.get('start_time')
    deadline_str = data.get('deadline')
    end_date_str = data.get('end_time')

    if not filename:
        return jsonify({"error": "Nama Dokumen harus diisi."}), 400
    
    if not start_date_str:
        return jsonify({"error": "Tanggal Mulai harus diisi."}), 400

    try:
        start_time = parse_flexible_date(start_date_str)
        deadline_date = parse_flexible_date(deadline_str)
        deadline = deadline_date.date() if deadline_date else None
        end_time = parse_flexible_date(end_date_str)
        status = calculate_task_status(start_time, deadline, end_time)
        
        new_log = AnalysisLog(
            user_id=current_user.id,
            document_type=data.get('document_type'),
            filename=filename,
            feature_type=feature_type,
            start_time=start_time,
            deadline=deadline,
            end_time=end_time,
            status=status
        )
        db.session.add(new_log)
        db.session.commit()
        
        return jsonify({"status": "success", "message": "Tugas berhasil ditambahkan."}), 201

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menambah tugas manual: {e}")
        return jsonify({"error": "Gagal menambah tugas."}), 500   

@app.route('/api/get_analysis_logs', methods=['GET'])
@login_required
def api_get_analysis_logs():
    """Mengambil semua log analisis untuk user saat ini."""
    try:
        logs = AnalysisLog.query.filter_by(user_id=current_user.id).order_by(AnalysisLog.start_time.desc()).all()
        logs_data = []
        today_date = datetime.date.today() 

        for log in logs:
            deadline_str = ''
            current_status = log.status
            log_deadline_date = None
            if log.deadline:
                if isinstance(log.deadline, datetime.datetime):
                    log_deadline_date = log.deadline.date()
                else:
                    log_deadline_date = log.deadline 
                
                deadline_str = log.deadline.strftime('%d %B %Y') 
                
                if current_status != 'done' and log_deadline_date < today_date:
                    current_status = 'overdue'

            def to_iso(dt):
                if not dt: return ""
                if isinstance(dt, datetime.datetime):
                    return dt.isoformat()[:16]
                return dt.strftime('%Y-%m-%dT00:00')

            logs_data.append({
                "id": log.id,
                "filename": log.filename,
                "feature_type": log.feature_type,
                "start_time": log.start_time.strftime('%d %B %Y') if log.start_time else "",
                "deadline": deadline_str,
                "end_time": log.end_time.strftime('%d %B %Y') if log.end_time else "-",
                "start_time_iso": to_iso(log.start_time),
                "deadline_iso": to_iso(log.deadline),
                "end_time_iso": to_iso(log.end_time),
                
                "status": current_status
            })
            
        return jsonify(logs_data), 200

    except Exception as e:
        print(f"!!! ERROR di api_get_analysis_logs: {e}")
        return jsonify({"error": f"Gagal mengambil log: {str(e)}"}), 500

@app.route('/api/edit_task/<int:log_id>', methods=['POST'])
@login_required
def api_edit_task(log_id):
    """Mengedit tugas yang sudah ada."""
    log = AnalysisLog.query.get(log_id)
    if not log or log.user_id != current_user.id:
        return jsonify({"error": "Tugas tidak ditemukan atau akses ditolak."}), 404

    data = request.get_json()

    log.filename = data.get('filename', log.filename)
    log.feature_type = data.get('feature_type', log.feature_type)
    
    try:
        def parse_dt(dt_str):
            if not dt_str: return None
            try: return datetime.datetime.strptime(dt_str, '%Y-%m-%dT%H:%M')
            except: return None

        if 'start_time' in data: 
            new_start = parse_dt(data['start_time'])
            if new_start: log.start_time = new_start
            
        if 'deadline' in data:
            new_dead = parse_dt(data['deadline'])
            log.deadline = new_dead.date() if new_dead else None 
            
        if 'end_time' in data:
            log.end_time = parse_dt(data['end_time'])

        log.status = calculate_task_status(log.start_time, log.deadline, log.end_time)
        
        db.session.commit()
        return jsonify({"status": "success", "message": "Tugas berhasil diperbarui."}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error edit task: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete_task/<int:log_id>', methods=['DELETE'])
@login_required
def api_delete_task(log_id):
    log = AnalysisLog.query.get(log_id)
    if not log or log.user_id != current_user.id:
        return jsonify({"error": "Tugas tidak ditemukan."}), 404

    try:
        db.session.delete(log)
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/monitoring_ams')
@login_required
def monitoring_ams_page():
    """Menampilkan halaman Monitoring AMS."""
    return render_template('monitoring_ams.html', 
                           username=current_user.username, 
                           label=current_user.label)

@app.route('/api/send_message', methods=['POST'])
@login_required
def api_send_message():
    """Mengirim pesan ke user lain tanpa dukungan lampiran."""
    recipient_id = request.form.get('recipient_id')
    subject = request.form.get('subject')
    body = request.form.get('body')

    if not recipient_id or not subject:
        return jsonify({"error": "Penerima dan Subjek harus diisi."}), 400

    try:
        attachment_path = None
        original_filename = None
        new_message = Message(
            sender_id=current_user.id,
            recipient_id=int(recipient_id), 
            subject=subject,
            body=body,
            attachment_path=attachment_path,
            original_filename=original_filename
        )
        db.session.add(new_message)
        db.session.commit()

        return jsonify({"status": "success", "message": "Pesan berhasil dikirim."}), 201

    except ValueError:
        db.session.rollback()
        return jsonify({"error": "ID penerima tidak valid."}), 400
    except Exception as e:
        db.session.rollback()
        print(f"Error saat mengirim pesan: {e}")
        return jsonify({"error": "Gagal mengirim pesan."}), 500

@app.route('/api/delete_message', methods=['POST'])
@login_required
def api_delete_message():
    data = request.json
    message_id = data.get('message_id')

    if not message_id:
        return jsonify({"error": "Message ID diperlukan."}), 400

    message = Message.query.get(message_id)

    if not message:
        return jsonify({"error": "Pesan tidak ditemukan."}), 404

    if message.sender_id != current_user.id and message.recipient_id != current_user.id:
        return jsonify({"error": "Anda tidak memiliki izin untuk menghapus pesan ini."}), 403

    try:
        db.session.delete(message)
        db.session.commit()
        return jsonify({"status": "success", "message": "Pesan berhasil dihapus secara mutual."}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting message: {e}")
        return jsonify({"error": "Terjadi kesalahan saat menghapus pesan di database."}), 500

@app.route('/api/generate_email_body', methods=['POST']) 
@login_required
def api_generate_email_body():
    """Menggunakan AI untuk menghasilkan draf isi email dari prompt."""
    data = request.json
    prompt = data.get('prompt')

    if not prompt:
        return jsonify({"error": "Prompt tidak boleh kosong."}), 400

    full_prompt = f"""
    Anda adalah asisten yang cerdas dan profesional dalam konteks kantor audit internal. 
    Tugas Anda adalah membuat draf isi email (body) yang formal, sopan, dan efektif dalam Bahasa Indonesia. 
    Format respons harus berupa body email siap pakai, tanpa subjek, tanpa penerima, dan tanpa salam penutup formal (cukup body utama).

    Gunakan konteks ini:
    - Pengirim adalah auditor dari SKAI IFG.
    - Nada harus tegas namun kolaboratif.
    - Pastikan semua permintaan/informasi jelas dan spesifik.
    - Mulai langsung dengan paragraf pembuka.
    - Isi body email selalu dimulai dengan "Dear (Penerima)"
    - Akhiri dengan kata terima kasih atas kerjasamanya dan improvisasi untuk menutup email
    - Ketika hasil draftnya sudah muncul di "Isi Email", hilangi kata-kata "Tentu, berikut adalah draf isi email yang sesuai dengan permintaan Anda."

    Berikut adalah instruksi/prompt dari pengguna:
    ---
    {prompt}
    ---

    Berikan hasil draf body email Anda.
    """
    
    try:
        ai_model = genai.GenerativeModel('gemini-2.5-flash') 
        track_gemini_usage()
        response = ai_model.generate_content(full_prompt)
        clean_body = response.text.strip().replace('**', '').replace('*', '')

        return jsonify({"status": "success", "body": clean_body}), 200
        
    except Exception as e:
        print(f"Error AI generating email: {e}")
        return jsonify({"error": "Gagal menghubungi layanan AI. Coba lagi nanti."}), 500

@app.route('/api/rename_monitoring_session', methods=['POST'])
@login_required
def api_rename_monitoring_session():
    """Mengubah nama sesi monitoring."""
    data = request.json
    old_name = data.get('old_name')
    new_name = data.get('new_name')

    if not old_name or not new_name:
        return jsonify({"error": "Nama lama dan baru harus diisi."}), 400

    if old_name == new_name:
        return jsonify({"message": "Nama tidak berubah."}), 200

    try:
        existing = AmsMonitoring.query.filter_by(user_id=current_user.id, session_name=new_name).first()
        if existing:
            return jsonify({"error": "Nama sesi baru sudah digunakan. Pilih nama lain."}), 400
        AmsMonitoring.query.filter_by(user_id=current_user.id, session_name=old_name).update(
            {AmsMonitoring.session_name: new_name}
        )

        SharedMonitoring.query.filter_by(owner_id=current_user.id, session_name=old_name).update(
            {SharedMonitoring.session_name: new_name}
        )

        db.session.commit()
        return jsonify({"status": "success", "message": "Nama sesi berhasil diubah."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error rename session: {e}")
        return jsonify({"error": "Gagal mengubah nama sesi."}), 500

@app.route('/api/delete_monitoring_session/<session_name>', methods=['DELETE'])
@login_required
def api_delete_monitoring_session(session_name):
    """Menghapus seluruh sesi monitoring dan semua auditee di dalamnya."""
    try:
        deleted_count = AmsMonitoring.query.filter_by(
            user_id=current_user.id, 
            session_name=session_name
        ).delete()
        
        db.session.commit()
        
        if deleted_count == 0:
            return jsonify({"error": "Sesi tidak ditemukan atau sudah kosong."}), 404

        return jsonify({
            "status": "success", 
            "message": f"Sesi '{session_name}' dan {deleted_count} data auditee berhasil dihapus."
        }), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus sesi monitoring: {e}")
        return jsonify({"error": "Gagal menghapus sesi."}), 500

@app.route('/api/get_messages', methods=['POST'])
@login_required
def api_get_messages():
    """Mengambil daftar pesan (inbox atau sent)."""
    data = request.json
    msg_type = data.get('type') 

    if msg_type == 'inbox':
        messages = Message.query.filter_by(recipient_id=current_user.id).order_by(Message.timestamp.desc()).all()
        
        messages_data = []
        for msg in messages:
             messages_data.append({
                 "id": msg.id,
                 "other_user": msg.sender.fullname if msg.sender else "Unknown Sender",
                 "subject": msg.subject,
                 "body": msg.body,
                 "timestamp": msg.timestamp.isoformat(), 
                 "is_read": msg.is_read,
                 "has_attachment": bool(msg.original_filename)
             })
    elif msg_type == 'sent':
        messages = Message.query.filter_by(sender_id=current_user.id).order_by(Message.timestamp.desc()).all()
        
        messages_data = []
        for msg in messages:
             messages_data.append({
                 "id": msg.id,
                 "other_user": msg.recipient.fullname if msg.recipient else "Unknown Recipient",
                 "subject": msg.subject,
                 "body": msg.body,
                 "timestamp": msg.timestamp.isoformat(), 
                 "is_read": msg.is_read,
                 "has_attachment": bool(msg.original_filename)
             })
    else:
        return jsonify({"error": "Tipe pesan tidak valid."}), 400

    return jsonify(messages_data), 200

@app.route('/ams_reminder')
@login_required
def ams_reminder_page():
    """Menampilkan halaman Reminder AMS."""
    return render_template('ams_reminder.html', 
                           username=current_user.username, 
                           label=current_user.label)

@app.route('/api/tl_tidak_setuju/delete_all', methods=['DELETE'])
@login_required
def api_delete_all_tl_tidak_setuju():
    try:
        num_deleted = db.session.query(AmsTlTidakSetuju).filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil menghapus {num_deleted} data."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Gagal menghapus data: {str(e)}"}), 500

@app.route('/api/download_message_attachment/<int:message_id>')
@login_required
def api_download_message_attachment(message_id):
    """Mengunduh lampiran pesan."""
    message = Message.query.get(message_id)
    if not message or (message.sender_id != current_user.id and message.recipient_id != current_user.id):
        return jsonify({"error": "Akses ditolak."}), 403

    if not message.attachment_path or not os.path.exists(message.attachment_path):
        return jsonify({"error": "File tidak ditemukan."}), 404

    return send_file(message.attachment_path, as_attachment=True, download_name=message.original_filename)

@app.route('/api/update_linked_pic_skai', methods=['POST'])
@login_required
def api_update_linked_pic_skai():
    """Update kolom PIC SKAI untuk beberapa baris sekaligus."""
    data = request.json
    updates = data.get('updates', []) 

    if not updates:
        return jsonify({"error": "Tidak ada data update."}), 400

    try:
        count = 0
        for item in updates:
            row_id = item.get('id')
            new_pic = item.get('pic_skai')
            
            row = AmsTemuanRow.query.get(row_id)
            if row:
                row.pic_skai = new_pic
                row.last_modified_by = current_user.username
                row.last_modified_at = datetime.datetime.now()
                count += 1
        
        db.session.commit()
        return jsonify({"status": "success", "message": f"Berhasil update {count} data PIC SKAI."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/get_unread_count', methods=['GET'])
@login_required
def api_get_unread_count():
    """Menghitung jumlah pesan yang belum dibaca untuk user saat ini."""
    try:
        count = Message.query.filter_by(recipient_id=current_user.id, is_read=False).count()
        return jsonify({"count": count}), 200
    except Exception as e:
        print(f"Error saat menghitung pesan belum dibaca: {e}")
        return jsonify({"error": "Gagal menghitung pesan."}), 500

@app.route('/api/mark_message_read/<int:message_id>', methods=['POST'])
@login_required
def api_mark_message_read(message_id):
    """Menandai pesan sebagai sudah dibaca."""
    message = Message.query.get(message_id)
    if not message or message.recipient_id != current_user.id:
        return jsonify({"error": "Pesan tidak ditemukan atau akses ditolak."}), 404

    try:
        message.is_read = True
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menandai pesan sebagai dibaca: {e}")
        return jsonify({"error": "Gagal memperbarui status pesan."}), 500

@app.route('/api/update_reminder_response/<int:id>', methods=['POST'])
@login_required
def api_update_reminder_response(id):
    """API untuk menyimpan status checkbox 'Sudah direspon' DAN 'Sudah diremind'"""
    data = request.json
    reminder = AmsReminder.query.get(id)
    
    if not reminder:
        return jsonify({"error": "Reminder tidak ditemukan."}), 404
    is_owner = (reminder.user_id == current_user.id)
    
    is_pic_skai = (current_user.username in reminder.pic_skai) if reminder.pic_skai else False
    is_pic_reminder = (current_user.username in reminder.pic_reminder) if reminder.pic_reminder else False
    is_pic_auditee = (current_user.username in reminder.pic_auditee) if reminder.pic_auditee else False
    has_access = is_owner or is_pic_skai or is_pic_reminder or is_pic_auditee

    if not has_access:
        return jsonify({"error": "Akses ditolak. Anda tidak terdaftar dalam reminder ini."}), 403
    
    try:
        if 'is_responded' in data:
            reminder.is_responded = data.get('is_responded')

        if 'is_reminded' in data:
            if is_owner or is_pic_skai or is_pic_reminder: 
                reminder.is_reminded = data.get('is_reminded')
            else:
                return jsonify({"error": "Hanya PIC SKAI/Reminder yang boleh mengubah status Remind."}), 403

        db.session.commit()
        return jsonify({"status": "success", "message": "Status berhasil diperbarui."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/save_reminder_temuan_link', methods=['POST'])
@login_required
def api_save_reminder_temuan_link():
    """Menyimpan relasi antara Reminder dan Temuan Row ID"""
    data = request.json
    reminder_id = data.get('reminder_id')
    temuan_ids = data.get('temuan_ids')
    reminder = AmsReminder.query.get(reminder_id)
    is_owner = (reminder.user_id == current_user.id)
    is_pic_skai = (current_user.username in reminder.pic_skai) if reminder.pic_skai else False
    is_pic_reminder = (current_user.username in reminder.pic_reminder) if reminder.pic_reminder else False

    if not reminder or (not is_owner and not is_pic_skai and not is_pic_reminder):
        return jsonify({"error": "Akses ditolak."}), 403

    try:
        if temuan_ids and len(temuan_ids) > 0:
            reminder.linked_temuan_ids = ",".join(map(str, temuan_ids))
            reminder.linked_by = current_user.username 
        else:
            reminder.linked_temuan_ids = None
            reminder.linked_by = None
            
        db.session.commit()
        return jsonify({"status": "success", "message": "Temuan berhasil ditautkan."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500  
    
@app.route('/api/unlink_reminder_temuan', methods=['POST'])
@login_required
def api_unlink_reminder_temuan():
    """Menghapus tautan temuan dari reminder"""
    data = request.json
    reminder_id = data.get('reminder_id')

    if not reminder_id:
        return jsonify({"error": "ID Reminder tidak valid."}), 400

    reminder = AmsReminder.query.get(reminder_id)
    is_owner = (reminder.user_id == current_user.id)
    is_pic_skai = (current_user.username in reminder.pic_skai) if reminder.pic_skai else False
    is_pic_reminder = (current_user.username in reminder.pic_reminder) if reminder.pic_reminder else False

    if not reminder or (not is_owner and not is_pic_skai and not is_pic_reminder):
        return jsonify({"error": "Akses ditolak."}), 403

    try:
        reminder.linked_temuan_ids = None 
        reminder.linked_by = None
        db.session.commit()
        return jsonify({"status": "success", "message": "Semua temuan berhasil dihapus."}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/view_linked_temuan/<int:reminder_id>')
@login_required
def view_linked_temuan(reminder_id):
    reminder = AmsReminder.query.get_or_404(reminder_id)
    
    linked_data = []
    if reminder.linked_temuan_ids:
        ids_list = [int(x) for x in reminder.linked_temuan_ids.split(',') if x.isdigit()]
        all_rows = AmsTemuanRow.query.filter(AmsTemuanRow.id.in_(ids_list)).all()
        is_viewer_admin = (
            reminder.user_id == current_user.id or \
            reminder.linked_by == current_user.username or \
            (reminder.pic_reminder and current_user.username in reminder.pic_reminder)
        )

        if is_viewer_admin:
            linked_data = all_rows
        else:
            filtered_rows = []
            for r in all_rows:
                if r.pic_skai and current_user.username in r.pic_skai:
                    filtered_rows.append(r)
            
            linked_data = filtered_rows

    return render_template('view_linked_temuan.html', reminder=reminder, rows=linked_data, username=current_user.username)

@app.route('/view_session_full/<int:session_id>')
@login_required
def view_session_full(session_id):
    session = AmsTemuanSession.query.get_or_404(session_id)
    is_owner = (session.user_id == current_user.id)
    is_shared = False
    if not is_owner:
        share_check = SharedTemuanSession.query.filter_by(
            session_id=session_id, 
            shared_with_id=current_user.id
        ).first()
        if share_check:
            is_shared = True
    
    if not is_owner and not is_shared:
        flash("Akses ditolak ke sesi ini.")
        return redirect(url_for('ams_temuan_page'))

    rows = AmsTemuanRow.query.filter_by(session_id=session_id).all()
    return render_template('view_session_full.html', session=session, rows=rows, username=current_user.username)

@app.route('/assign_pic_page/<int:reminder_id>')
@login_required
def assign_pic_page(reminder_id):
    """Halaman khusus untuk Assign PIC SKAI pada temuan yang sudah dilink."""
    reminder = AmsReminder.query.get_or_404(reminder_id)
    
    linked_data = []
    if reminder.linked_temuan_ids:
        ids_list = [int(x) for x in reminder.linked_temuan_ids.split(',') if x.isdigit()]
        linked_data = AmsTemuanRow.query.filter(AmsTemuanRow.id.in_(ids_list)).all()
    users = User.query.all() 
    
    return render_template('assign_pic_page.html', 
                           reminder=reminder, 
                           rows=linked_data, 
                           users=users,
                           username=current_user.username)

@app.route('/api/gemini-status', methods=['GET'])
def get_gemini_status():
    global gemini_usage_count, gemini_last_reset
    
    now = datetime.datetime.now()
    time_diff = (now - gemini_last_reset).total_seconds()
    
    current_usage = gemini_usage_count
    reset_seconds = 60 - int(time_diff)
    
    if time_diff >= 60:
        current_usage = 0
        reset_seconds = 60
        gemini_last_reset = now 
        gemini_usage_count = 0

    return jsonify({
        "used": current_usage,
        "limit": RPM_LIMIT,
        "reset_seconds": max(0, reset_seconds) 
    })

@app.route('/monitoring_dashboard')
@login_required
def monitoring_dashboard_page():
    return render_template('monitoring_dashboard.html', username=current_user.username, label=current_user.label)

@app.route('/api/get_dashboard_trend_data', methods=['GET'])
@login_required
def api_get_dashboard_trend_data():
    try:
        all_data = AmsMonitoring.query.filter(
            or_(AmsMonitoring.user_id == current_user.id,
                AmsMonitoring.session_name.in_(
                    db.session.query(SharedMonitoring.session_name)
                    .filter(SharedMonitoring.shared_with_id == current_user.id)
                ))
        ).all()

        processed_data = []

        # Dictionary Nama Bulan -> Angka
        month_map = {
            'januari': '01', 'februari': '02', 'maret': '03', 'april': '04', 
            'mei': '05', 'juni': '06', 'juli': '07', 'agustus': '08', 
            'september': '09', 'oktober': '10', 'november': '11', 'desember': '12',
            'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04', 'jun': '06', 
            'jul': '07', 'aug': '08', 'agu': '08', 'sep': '09', 'okt': '10', 
            'nov': '11', 'dec': '12', 'des': '12'
        }

        # Nama bulan untuk label yang cantik
        month_labels = {
            '01': 'Januari', '02': 'Februari', '03': 'Maret', '04': 'April',
            '05': 'Mei', '06': 'Juni', '07': 'Juli', '08': 'Agustus',
            '09': 'September', '10': 'Oktober', '11': 'November', '12': 'Desember'
        }

        for item in all_data:
            # 1. Bersihkan nama sesi & jadikan huruf kecil semua
            original_name = item.session_name
            clean_name = original_name.lower().strip()
            
            # --- LOGIKA "BRUTE FORCE" DETEKSI TANGGAL ---
            detected_month = None
            detected_year = None
            
            # Cari Tahun (4 digit angka, misal 2024, 2025, 2026)
            year_match = re.search(r'\d{4}', clean_name)
            if year_match:
                detected_year = year_match.group(0)

            # Cari Bulan (Looping cek satu per satu)
            for m_name, m_code in month_map.items():
                # Cek apakah kata bulan ada di dalam nama sesi (misal: "januari" ada di "audit internal - januari 2026")
                if m_name in clean_name:
                    detected_month = m_code
                    break # Berhenti jika sudah ketemu
            
            # --- PENENTUAN PERIODE FINAL ---
            if detected_month and detected_year:
                # KASUS SUKSES: Nama sesi mengandung Bulan dan Tahun
                periode_iso = f"{detected_year}-{detected_month}"
                label_bulan = month_labels.get(detected_month, 'Unknown')
                periode_label = f"{label_bulan} {detected_year}"
                
                # Bersihkan Kategori Audit (Hapus bulan & tahun dari nama)
                # Contoh: "Audit Internal - Januari 2026" -> menjadi "Audit Internal"
                audit_category = re.sub(r'\d{4}', '', original_name) # Hapus tahun
                audit_category = re.sub(r'(?i)' + label_bulan, '', audit_category) # Hapus nama bulan
                audit_category = re.sub(r'[\-\_\/,\.]+$', '', audit_category.strip()).strip() # Hapus simbol sisa di akhir
                
            else:
                # KASUS GAGAL: Gunakan tanggal dari Database sebagai cadangan
                if item.periode:
                    periode_iso = item.periode.strftime('%Y-%m')
                    periode_label = item.periode.strftime('%B %Y')
                else:
                    periode_iso = "0000-00"
                    periode_label = "Tanpa Periode"
                audit_category = original_name # Kategori tetap nama asli

            processed_data.append({
                "audit_category": audit_category, # Kategori bersih (untuk filter)
                "session_name": original_name,
                "periode": periode_iso,       # PENTING: Ini kunci pemisah (YYYY-MM)
                "periode_label": periode_label, 
                "auditee": item.auditee,
                "selesai": item.selesai,
                "tidak_selesai": item.tidak_selesai,
                "todo": item.todo,
                "belum_sesuai": item.belum_sesuai,
                "belum_tl": item.belum_tl,
                "tdd": item.tdd,
                "type": item.monitoring_type
            })
        
        # Urutkan berdasarkan YYYY-MM
        processed_data.sort(key=lambda x: x['periode'])
        
        return jsonify(processed_data), 200

    except Exception as e:
        print(f"Error Dashboard Data: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate_dashboard_insight', methods=['POST'])
@login_required
def api_generate_dashboard_insight():
    data = request.json
    audit_category = data.get('category', 'Audit Umum')
    chart_data = data.get('chart_data', {})
    
    # Validasi data
    if not chart_data or not chart_data.get('labels'):
        return jsonify({"error": "Data grafik tidak tersedia untuk dianalisis."}), 400

    # Susun prompt yang TERSTRUKTUR PER SECTION
    prompt = f"""
    Anda adalah Chief Audit Executive (CAE) yang ahli dalam menyajikan data.
    Tugas Anda adalah memberikan interpretasi data audit yang **singkat, padat, dan estetik** berdasarkan data berikut:
    
    - Kategori Audit: {audit_category}
    - Periode Waktu: {', '.join(chart_data['labels'])}
    - Tren 'Selesai': {chart_data['selesai']}
    - Tren 'Belum Sesuai/BJT': {chart_data['bjt']}
    - Tren 'Outstanding/Belum TL': {chart_data['outstanding']}
    
    Instruksi Output (Wajib Format HTML):
    Sajikan analisis Anda dalam 3 bagian terpisah. Jangan gunakan paragraf panjang, gunakan **poin-poin (bullet points)** saja.
    Jangan gunakan Markdown (```), langsung tag HTML.

    Format Struktur Jawaban:
    
    <div style="margin-bottom: 15px; text-align: justify;">
        <h4 style="color: #2c3e50; margin-bottom: 8px; border-bottom: 2px solid #4CAF50; display: inline-block;">1. Analisis Tren (Line Chart)</h4>
        <ul style="margin-top: 5px; color: #444;">
            <li>Analisis pergerakan grafik: Apakah tren penyelesaian naik/turun secara signifikan antar periode?</li>
            <li>Identifikasi anomali: Adakah bulan tertentu yang memiliki lonjakan Outstanding yang mengkhawatirkan?</li>
        </ul>
    </div>

    <div style="margin-bottom: 15px; text-align: justify;">
        <h4 style="color: #2c3e50; margin-bottom: 8px; border-bottom: 2px solid #FF9800; display: inline-block;">2. Komposisi Status (Overview)</h4>
        <ul style="margin-top: 5px; color: #444;">
            <li>Secara akumulatif, status mana yang mendominasi porsi grafik? (Selesai vs Outstanding).</li>
            <li>Berikan penilaian singkat: Apakah komposisi ini menunjukkan kinerja audit yang Sehat, Waspada, atau Kritis?</li>
        </ul>
    </div>

    <div>
        <h4 style="color: #2c3e50; margin-bottom: 8px; border-bottom: 2px solid #F44336; display: inline-block; text-align: justify;">3. Rekomendasi Strategis</h4>
        <ul style="margin-top: 5px; color: #444;">
            <li>Saran taktis 1: Apa yang harus dilakukan PIC/Auditor segera?</li>
            <li>Saran taktis 2: Fokus perbaikan untuk periode berikutnya.</li>
        </ul>
    </div>

    Gunakan Bahasa Indonesia yang profesional, tegas, namun mudah dipahami.
    """

    try:
        model = genai.GenerativeModel('gemini-2.5-flash') 
        track_gemini_usage()
        response = model.generate_content(prompt)
        
        clean_text = response.text.replace('```html', '').replace('```', '')
        
        return jsonify({"status": "success", "analysis": clean_text}), 200
    except Exception as e:
        print(f"Error AI Insight: {e}")
        return jsonify({"error": "Gagal menghasilkan analisis AI."}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)